"""
Document Generation Pipeline for the Corporate Law Document Generator.

Each document type is defined as a dict with:
  label, description, fields (list of form field defs), system_prompt.

The pipeline:
  1. Retrieve style examples from ChromaDB (where document_type = type)
  2. Optionally fetch reference data from SEC EDGAR
  3. Build system + user prompts
  4. Call LLMBackend.generate_document()
  5. Convert result to .docx bytes
"""

import io
import re
import difflib
import zipfile
import xml.etree.ElementTree as ET
from datetime import date

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from llm_backend import LLMBackend
from extended_document_types import ADDITIONAL_DOCUMENT_TYPES


# ======================================================================
# Document type definitions
# ======================================================================

DOCUMENT_TYPES: dict[str, dict] = {
    "contract": {
        "label": "Contract / Agreement",
        "description": "Generate a contract or agreement between parties.",
        "fields": [
            {"key": "party_a", "label": "Party A (Full Legal Name)", "type": "text"},
            {"key": "party_b", "label": "Party B (Full Legal Name)", "type": "text"},
            {"key": "effective_date", "label": "Effective Date", "type": "date"},
            {"key": "term", "label": "Term / Duration", "type": "text",
             "placeholder": "e.g. 12 months, perpetual"},
            {"key": "subject_matter", "label": "Subject Matter", "type": "textarea",
             "placeholder": "Describe the purpose and scope of the agreement"},
            {"key": "key_terms", "label": "Key Terms & Conditions", "type": "textarea",
             "placeholder": "Payment terms, deliverables, obligations, etc."},
            {"key": "governing_law", "label": "Governing Law (State/Jurisdiction)", "type": "text",
             "placeholder": "e.g. State of Delaware"},
            {"key": "special_provisions", "label": "Special Provisions (optional)", "type": "textarea"},
        ],
        "system_prompt": (
            "You are a corporate law document drafting assistant. "
            "Generate a professional, well-structured contract/agreement. "
            "Use formal legal language with numbered sections and sub-sections. "
            "Include standard boilerplate clauses (entire agreement, severability, "
            "notices, counterparts) unless the user specifies otherwise. "
            "Do NOT provide legal advice — this is a draft for attorney review."
        ),
    },
    "memo": {
        "label": "Legal Memorandum",
        "description": "Generate an internal legal memorandum.",
        "fields": [
            {"key": "to", "label": "To", "type": "text"},
            {"key": "from_field", "label": "From", "type": "text"},
            {"key": "date_field", "label": "Date", "type": "date"},
            {"key": "re", "label": "Re (Subject)", "type": "text"},
            {"key": "questions_presented", "label": "Questions Presented", "type": "textarea",
             "placeholder": "List the legal questions to be addressed"},
            {"key": "facts", "label": "Relevant Facts", "type": "textarea"},
            {"key": "jurisdiction", "label": "Jurisdiction", "type": "text",
             "placeholder": "e.g. Federal, State of New York"},
        ],
        "system_prompt": (
            "You are a corporate law document drafting assistant. "
            "Generate a professional internal legal memorandum. "
            "Follow the standard memo format: Header block (TO/FROM/DATE/RE), "
            "Questions Presented, Brief Answer, Discussion, Conclusion. "
            "Cite relevant legal principles where appropriate. "
            "Do NOT provide legal advice — this is a draft for attorney review."
        ),
    },
    "brief": {
        "label": "Legal Brief",
        "description": "Generate a legal brief or motion.",
        "fields": [
            {"key": "court", "label": "Court", "type": "text",
             "placeholder": "e.g. United States District Court, Southern District of New York"},
            {"key": "case_caption", "label": "Case Caption", "type": "text",
             "placeholder": "e.g. Smith v. Jones Corp."},
            {"key": "case_number", "label": "Case Number", "type": "text"},
            {"key": "brief_type", "label": "Brief Type", "type": "text",
             "placeholder": "e.g. Motion to Dismiss, Summary Judgment, Opposition"},
            {"key": "argument_summary", "label": "Argument Summary", "type": "textarea",
             "placeholder": "Summarize the key arguments to be made"},
            {"key": "facts", "label": "Statement of Facts", "type": "textarea"},
            {"key": "legal_standard", "label": "Legal Standard", "type": "textarea",
             "placeholder": "Applicable legal standard or rule"},
        ],
        "system_prompt": (
            "You are a corporate law document drafting assistant. "
            "Generate a professional legal brief or motion. "
            "Follow standard brief format: Caption, Introduction, Statement of Facts, "
            "Argument (with numbered points and sub-points), Conclusion, Signature Block. "
            "Use formal legal writing style with proper citations (Bluebook format). "
            "Do NOT provide legal advice — this is a draft for attorney review."
        ),
    },
    "filing": {
        "label": "Corporate Filing",
        "description": "Generate a corporate filing document (articles, annual report, etc.).",
        "fields": [
            {"key": "filing_type", "label": "Filing Type", "type": "text",
             "placeholder": "e.g. Articles of Incorporation, Annual Report, Certificate of Amendment"},
            {"key": "entity_name", "label": "Entity Name", "type": "text"},
            {"key": "state", "label": "State of Formation / Filing", "type": "text"},
            {"key": "details", "label": "Filing Details", "type": "textarea",
             "placeholder": "Purpose, authorized shares, registered agent, amendments, etc."},
            {"key": "sec_cik", "label": "SEC CIK (optional, for reference data)", "type": "text",
             "placeholder": "10-digit CIK number for SEC EDGAR lookup"},
        ],
        "system_prompt": (
            "You are a corporate law document drafting assistant. "
            "Generate a professional corporate filing document. "
            "Use the correct format for the specified filing type and jurisdiction. "
            "Include all required statutory sections and language. "
            "Do NOT provide legal advice — this is a draft for attorney review."
        ),
    },
}


# ======================================================================
# Generator
# ======================================================================


class DocumentGenerator:
    """End-to-end generation pipeline: examples -> ref data -> prompt -> LLM -> text."""

    def __init__(self, llm: LLMBackend, chroma_collection=None, knowledge_db=None):
        self.llm = llm
        self.collection = chroma_collection
        self.knowledge_db = knowledge_db  # Optional: for learned templates

    # -- Document Analysis (for Mimic workflow) --------------------------

    def analyze_document(self, document_text: str) -> dict:
        """
        Analyze an uploaded document to extract structure, fields, and style.
        Returns a dict with document_type, fields, structure, tone, etc.
        """
        analysis_prompt = """
You are a legal document analysis expert. Analyze the provided document and extract the following information in JSON format:

{
    "document_type": "contract|memo|brief|filing|other",
    "document_subtype": "specific type, e.g., Independent Contractor Agreement, Lease, NDA, etc.",
    "structure": {
        "sections": ["list of main section headings"],
        "has_signature_block": true/false,
        "has_exhibits": true/false
    },
    "key_fields": {
        "field_name": "extracted value",
        ...
    },
    "tone": "formal|semi-formal|technical",
    "style_notes": "brief description of writing style and formatting patterns"
}

Extract all named parties, dates, amounts, terms, and other variable fields. For fields that appear to be templates or variables, extract the actual values if present.

Document to analyze:

"""

        messages = [
            {"role": "system", "content": "You are a legal document analysis expert. Always respond with valid JSON."},
            {"role": "user", "content": analysis_prompt + document_text[:8000]}  # Truncate to fit context
        ]

        try:
            response = self.llm.chat(messages, temperature=0.1, max_tokens=2048)
            # Try to extract JSON from the response
            import json
            # Find JSON in the response (might be wrapped in markdown code blocks)
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = response[json_start:json_end]
                analysis = json.loads(json_str)
                return analysis
            else:
                # Fallback if JSON parsing fails
                return {
                    "document_type": "contract",
                    "document_subtype": "Unknown",
                    "structure": {"sections": [], "has_signature_block": False, "has_exhibits": False},
                    "key_fields": {},
                    "tone": "formal",
                    "style_notes": "Unable to parse analysis"
                }
        except Exception as e:
            return {
                "document_type": "contract",
                "document_subtype": "Unknown",
                "structure": {"sections": [], "has_signature_block": False, "has_exhibits": False},
                "key_fields": {},
                "tone": "formal",
                "style_notes": f"Analysis error: {str(e)}"
            }

    def generate_from_template(
        self,
        reference_text: str,
        analysis: dict,
        user_edits: dict
    ) -> str:
        """
        Generate a new document that mimics the reference document's style
        but uses the user's edited field values.
        """
        system_prompt = f"""You are a legal document drafting assistant. Generate a new {analysis.get('document_subtype', 'document')} that:
1. Mimics the EXACT structure and formatting of the reference document
2. Uses the same tone and style: {analysis.get('tone', 'formal')}
3. Maintains the same section headings and organization
4. Replaces template values with the user's specified values

Style notes: {analysis.get('style_notes', 'Professional legal document')}

This is a draft for attorney review, not legal advice."""

        # Build the user prompt with reference and edits
        field_list = "\n".join([f"- {k}: {v}" for k, v in user_edits.items() if v])

        user_prompt = f"""Reference document structure and style:

{reference_text[:4000]}

User's values for this new document:

{field_list}

Generate a complete {analysis.get('document_subtype', 'document')} following the exact structure and style of the reference document, but using the user's values."""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]

        return self.llm.chat(messages, temperature=0.2, max_tokens=4096)

    # -- AI-Guided Workflow ----------------------------------------------

    def get_required_fields_for_type(self, document_type: str) -> list[dict]:
        """
        Ask the LLM what fields are typically needed for a document type.
        First checks learned templates from real documents, then falls back to predefined.
        Returns a list of field definitions.
        """
        # PRIORITY 1: Check if we have a learned template from real documents
        if self.knowledge_db:
            try:
                learned = self.knowledge_db.get_learned_template(document_type)
                if learned and learned.get("sample_count", 0) >= 3:  # Require at least 3 samples
                    # Use learned template
                    fields = learned["fields"]
                    # Add metadata to indicate these are learned
                    for field in fields:
                        field["learned"] = True
                        field["sample_count"] = learned["sample_count"]
                    return fields
            except Exception as e:
                # Fall through to hardcoded templates
                pass

        # PRIORITY 2: Enhanced predefined templates for common document types
        EXTENDED_DOCUMENT_TYPES = {
            "Independent Contractor Agreement": [
                {"key": "company_name", "label": "Company Name", "type": "text", "placeholder": "ABC Corp", "help": "Hiring company's legal name"},
                {"key": "contractor_name", "label": "Contractor Name", "type": "text", "placeholder": "John Doe", "help": "Independent contractor's full name"},
                {"key": "contractor_address", "label": "Contractor Address", "type": "text", "placeholder": "123 Main St, City, State", "help": "Contractor's address"},
                {"key": "services", "label": "Services Description", "type": "textarea", "placeholder": "Describe the services to be performed", "help": "Detailed description of work"},
                {"key": "compensation", "label": "Compensation Amount", "type": "text", "placeholder": "$5,000", "help": "Total payment or rate"},
                {"key": "payment_terms", "label": "Payment Terms", "type": "textarea", "placeholder": "Net 30 days, monthly invoicing", "help": "When and how payment is made"},
                {"key": "start_date", "label": "Start Date", "type": "date", "help": "Contract start date"},
                {"key": "end_date", "label": "End Date (if applicable)", "type": "date", "help": "Contract end date"},
                {"key": "ip_assignment", "label": "Intellectual Property Assignment", "type": "text", "placeholder": "Yes/No", "help": "Will contractor IP belong to company?"},
                {"key": "confidentiality", "label": "Confidentiality Requirements", "type": "textarea", "placeholder": "Describe confidentiality obligations", "help": "NDA and confidentiality terms"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of California", "help": "Which state's laws apply"},
                {"key": "termination", "label": "Termination Notice Period", "type": "text", "placeholder": "30 days", "help": "Notice required to terminate"},
            ],
            "NDA / Confidentiality Agreement": [
                {"key": "disclosing_party", "label": "Disclosing Party", "type": "text", "placeholder": "ABC Corp", "help": "Party sharing confidential information"},
                {"key": "receiving_party", "label": "Receiving Party", "type": "text", "placeholder": "XYZ LLC", "help": "Party receiving confidential information"},
                {"key": "purpose", "label": "Purpose", "type": "textarea", "placeholder": "Evaluation of potential business relationship", "help": "Reason for sharing information"},
                {"key": "effective_date", "label": "Effective Date", "type": "date", "help": "When the NDA becomes effective"},
                {"key": "term", "label": "Term / Duration", "type": "text", "placeholder": "2 years", "help": "How long confidentiality lasts"},
                {"key": "definition", "label": "Confidential Information Definition", "type": "textarea", "placeholder": "All technical, business, financial information...", "help": "What is considered confidential"},
                {"key": "exceptions", "label": "Exceptions (optional)", "type": "textarea", "placeholder": "Information already public, independently developed...", "help": "What is NOT confidential"},
                {"key": "return_obligation", "label": "Return of Materials", "type": "text", "placeholder": "Yes/No", "help": "Must materials be returned?"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of Delaware", "help": "Which jurisdiction applies"},
            ],
            "Employment Agreement": [
                {"key": "employer_name", "label": "Employer Name", "type": "text", "placeholder": "ABC Company Inc.", "help": "Employer's legal name"},
                {"key": "employee_name", "label": "Employee Name", "type": "text", "placeholder": "Jane Smith", "help": "Employee's full legal name"},
                {"key": "position", "label": "Position / Title", "type": "text", "placeholder": "Senior Software Engineer", "help": "Job title"},
                {"key": "start_date", "label": "Start Date", "type": "date", "help": "Employment start date"},
                {"key": "salary", "label": "Annual Salary", "type": "text", "placeholder": "$120,000", "help": "Annual compensation"},
                {"key": "benefits", "label": "Benefits", "type": "textarea", "placeholder": "Health insurance, 401(k), PTO...", "help": "Employee benefits package"},
                {"key": "duties", "label": "Duties & Responsibilities", "type": "textarea", "placeholder": "Primary job responsibilities", "help": "What the employee will do"},
                {"key": "work_location", "label": "Work Location", "type": "text", "placeholder": "San Francisco, CA / Remote", "help": "Where work is performed"},
                {"key": "at_will", "label": "At-Will Employment", "type": "text", "placeholder": "Yes/No", "help": "Is this at-will employment?"},
                {"key": "confidentiality", "label": "Confidentiality Clause", "type": "text", "placeholder": "Yes/No", "help": "Include confidentiality obligations?"},
                {"key": "non_compete", "label": "Non-Compete Clause", "type": "text", "placeholder": "Yes/No", "help": "Include non-compete?"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of California", "help": "Which state's laws apply"},
            ],
            "Operating Agreement / LLC": [
                {"key": "llc_name", "label": "LLC Name", "type": "text", "placeholder": "ABC Services LLC", "help": "Legal name of the LLC"},
                {"key": "state", "label": "State of Formation", "type": "text", "placeholder": "Delaware", "help": "State where LLC is formed"},
                {"key": "formation_date", "label": "Formation Date", "type": "date", "help": "Date LLC was formed"},
                {"key": "members", "label": "Members", "type": "textarea", "placeholder": "John Doe (50%), Jane Smith (50%)", "help": "List all members and ownership %"},
                {"key": "purpose", "label": "Business Purpose", "type": "textarea", "placeholder": "Provide consulting services", "help": "Purpose of the LLC"},
                {"key": "management", "label": "Management Type", "type": "text", "placeholder": "Member-managed / Manager-managed", "help": "How is the LLC managed?"},
                {"key": "capital", "label": "Initial Capital Contributions", "type": "textarea", "placeholder": "Member contributions and amounts", "help": "Initial capital each member contributes"},
                {"key": "distributions", "label": "Distribution Terms", "type": "textarea", "placeholder": "Pro-rata based on ownership", "help": "How profits are distributed"},
                {"key": "voting", "label": "Voting Rights", "type": "textarea", "placeholder": "Majority vote required", "help": "How decisions are made"},
                {"key": "transfer", "label": "Transfer Restrictions", "type": "textarea", "placeholder": "Right of first refusal", "help": "Can members sell their interest?"},
            ],
            "Letter / Correspondence": [
                {"key": "sender", "label": "Sender Name", "type": "text", "placeholder": "John Doe, Esq.", "help": "Who is sending the letter"},
                {"key": "recipient", "label": "Recipient Name", "type": "text", "placeholder": "Jane Smith", "help": "Who is receiving the letter"},
                {"key": "date", "label": "Date", "type": "date", "help": "Date of the letter"},
                {"key": "subject", "label": "Subject / Re:", "type": "text", "placeholder": "Regarding Contract Dispute", "help": "Subject line"},
                {"key": "body", "label": "Letter Content", "type": "textarea", "placeholder": "Main content of the letter", "help": "What do you want to communicate?"},
                {"key": "tone", "label": "Tone", "type": "text", "placeholder": "Formal / Friendly / Stern", "help": "Desired tone of the letter"},
                {"key": "action", "label": "Desired Action", "type": "textarea", "placeholder": "Request payment within 10 days", "help": "What action do you want the recipient to take?"},
            ],
            "Custom Document": [
                {"key": "document_description", "label": "Document Description", "type": "textarea", "placeholder": "Describe the type of document you need", "help": "What kind of document do you want to create?", "required": True},
                {"key": "parties", "label": "Parties Involved", "type": "textarea", "placeholder": "Who are the parties?", "help": "List all parties involved"},
                {"key": "key_terms", "label": "Key Terms & Details", "type": "textarea", "placeholder": "What are the important terms, conditions, or details?", "help": "Main content and requirements"},
                {"key": "special_provisions", "label": "Special Provisions (optional)", "type": "textarea", "placeholder": "Any specific clauses or provisions needed?", "help": "Additional requirements"},
            ],
            # ===== CORPORATE & BUSINESS =====
            "Shareholder Agreement": [
                {"key": "company_name", "label": "Company Name", "type": "text", "placeholder": "ABC Corporation", "help": "Legal name of the corporation"},
                {"key": "shareholders", "label": "Shareholders", "type": "textarea", "placeholder": "John Doe (40%), Jane Smith (35%), Bob Johnson (25%)", "help": "List all shareholders and their ownership percentages"},
                {"key": "formation_date", "label": "Company Formation Date", "type": "date", "help": "Date the company was incorporated"},
                {"key": "share_class", "label": "Share Class", "type": "text", "placeholder": "Common Stock, Class A", "help": "Type of shares covered"},
                {"key": "voting_rights", "label": "Voting Rights", "type": "textarea", "placeholder": "One vote per share, majority required for major decisions", "help": "How voting is structured"},
                {"key": "transfer_restrictions", "label": "Share Transfer Restrictions", "type": "textarea", "placeholder": "Right of first refusal, board approval required", "help": "Restrictions on selling or transferring shares"},
                {"key": "drag_along", "label": "Drag-Along Rights", "type": "text", "placeholder": "Yes/No", "help": "Can majority force minority to sell?"},
                {"key": "tag_along", "label": "Tag-Along Rights", "type": "text", "placeholder": "Yes/No", "help": "Can minority join a sale?"},
                {"key": "dividend_policy", "label": "Dividend Policy", "type": "textarea", "placeholder": "Distributed quarterly, pro-rata based on ownership", "help": "How profits are distributed"},
                {"key": "dispute_resolution", "label": "Dispute Resolution", "type": "textarea", "placeholder": "Mediation then arbitration", "help": "How disputes are resolved"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of Delaware", "help": "Which jurisdiction's laws apply"},
            ],
            "Partnership Agreement": [
                {"key": "partnership_name", "label": "Partnership Name", "type": "text", "placeholder": "Smith & Associates Partnership", "help": "Legal name of the partnership"},
                {"key": "partners", "label": "Partners", "type": "textarea", "placeholder": "Partner names and ownership percentages", "help": "List all partners with ownership stakes"},
                {"key": "business_purpose", "label": "Business Purpose", "type": "textarea", "placeholder": "Professional consulting services", "help": "Primary purpose of the partnership"},
                {"key": "start_date", "label": "Start Date", "type": "date", "help": "When the partnership begins"},
                {"key": "capital_contributions", "label": "Capital Contributions", "type": "textarea", "placeholder": "Initial capital and ongoing contribution requirements", "help": "How much each partner contributes"},
                {"key": "profit_distribution", "label": "Profit Distribution", "type": "textarea", "placeholder": "Distributed based on ownership percentage", "help": "How profits and losses are shared"},
                {"key": "decision_making", "label": "Decision Making Authority", "type": "textarea", "placeholder": "Unanimous for major decisions, majority for day-to-day", "help": "How decisions are made"},
                {"key": "management_duties", "label": "Management Duties", "type": "textarea", "placeholder": "Roles and responsibilities of each partner", "help": "Who does what"},
                {"key": "withdrawal", "label": "Withdrawal Terms", "type": "textarea", "placeholder": "90-day notice, buyout at fair market value", "help": "How a partner can exit"},
                {"key": "dissolution", "label": "Dissolution Terms", "type": "textarea", "placeholder": "Conditions for dissolution and asset distribution", "help": "How partnership can be dissolved"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of New York", "help": "Applicable jurisdiction"},
            ],
            "Stock Purchase Agreement": [
                {"key": "seller", "label": "Seller Name", "type": "text", "placeholder": "John Doe", "help": "Party selling the stock"},
                {"key": "buyer", "label": "Buyer Name", "type": "text", "placeholder": "Jane Smith", "help": "Party purchasing the stock"},
                {"key": "company_name", "label": "Company Name", "type": "text", "placeholder": "ABC Corporation", "help": "Company whose stock is being sold"},
                {"key": "shares_quantity", "label": "Number of Shares", "type": "number", "placeholder": "10000", "help": "Total shares being sold"},
                {"key": "share_class", "label": "Share Class", "type": "text", "placeholder": "Common Stock", "help": "Type of shares"},
                {"key": "purchase_price", "label": "Total Purchase Price", "type": "text", "placeholder": "$500,000", "help": "Total amount buyer will pay"},
                {"key": "price_per_share", "label": "Price Per Share", "type": "text", "placeholder": "$50.00", "help": "Price for each share"},
                {"key": "payment_terms", "label": "Payment Terms", "type": "textarea", "placeholder": "50% at closing, 50% in 30 days", "help": "How and when payment is made"},
                {"key": "closing_date", "label": "Closing Date", "type": "date", "help": "When the transaction closes"},
                {"key": "representations", "label": "Seller Representations", "type": "textarea", "placeholder": "Ownership, authority, no encumbrances", "help": "Seller's promises about the stock"},
                {"key": "warranties", "label": "Warranties", "type": "textarea", "placeholder": "Stock is validly issued, fully paid, non-assessable", "help": "Guarantees about the stock"},
                {"key": "conditions", "label": "Closing Conditions", "type": "textarea", "placeholder": "Due diligence completion, regulatory approvals", "help": "What must happen before closing"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of Delaware", "help": "Applicable jurisdiction"},
            ],
            "Asset Purchase Agreement": [
                {"key": "seller", "label": "Seller Name", "type": "text", "placeholder": "ABC Company Inc.", "help": "Party selling the assets"},
                {"key": "buyer", "label": "Buyer Name", "type": "text", "placeholder": "XYZ Acquisitions LLC", "help": "Party purchasing the assets"},
                {"key": "assets_description", "label": "Assets Being Sold", "type": "textarea", "placeholder": "Equipment, inventory, intellectual property, customer lists, goodwill", "help": "Detailed description of all assets"},
                {"key": "excluded_assets", "label": "Excluded Assets", "type": "textarea", "placeholder": "Cash, accounts receivable prior to closing", "help": "What is NOT being sold"},
                {"key": "purchase_price", "label": "Total Purchase Price", "type": "text", "placeholder": "$1,250,000", "help": "Total amount for all assets"},
                {"key": "allocation", "label": "Price Allocation", "type": "textarea", "placeholder": "Equipment: $500K, Inventory: $300K, IP: $400K, Goodwill: $50K", "help": "How price is allocated to each asset category"},
                {"key": "payment_terms", "label": "Payment Terms", "type": "textarea", "placeholder": "Cash at closing", "help": "How buyer will pay"},
                {"key": "assumed_liabilities", "label": "Assumed Liabilities", "type": "textarea", "placeholder": "Seller's lease obligations, certain contracts", "help": "Which liabilities buyer is taking on"},
                {"key": "closing_date", "label": "Closing Date", "type": "date", "help": "When transaction closes"},
                {"key": "transition_assistance", "label": "Transition Assistance", "type": "textarea", "placeholder": "Seller to provide 30 days training and support", "help": "Post-closing support from seller"},
                {"key": "non_compete", "label": "Non-Compete Terms", "type": "textarea", "placeholder": "Seller shall not compete for 3 years within 50 miles", "help": "Restrictions on seller's future business"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of California", "help": "Applicable jurisdiction"},
            ],
            "Merger Agreement": [
                {"key": "company_a", "label": "First Company", "type": "text", "placeholder": "ABC Corporation", "help": "Name of first merging company"},
                {"key": "company_b", "label": "Second Company", "type": "text", "placeholder": "XYZ Inc.", "help": "Name of second merging company"},
                {"key": "surviving_entity", "label": "Surviving Entity", "type": "text", "placeholder": "ABC Corporation", "help": "Which company will survive the merger"},
                {"key": "merger_type", "label": "Merger Type", "type": "text", "placeholder": "Forward triangular merger", "help": "Structure of the merger"},
                {"key": "exchange_ratio", "label": "Stock Exchange Ratio", "type": "text", "placeholder": "1.5 shares of ABC for each share of XYZ", "help": "How stocks will be converted"},
                {"key": "consideration", "label": "Total Consideration", "type": "text", "placeholder": "$10,000,000 cash + stock", "help": "Total value of the merger"},
                {"key": "closing_date", "label": "Expected Closing Date", "type": "date", "help": "Target date for merger completion"},
                {"key": "conditions_precedent", "label": "Closing Conditions", "type": "textarea", "placeholder": "Shareholder approval, regulatory clearances, due diligence", "help": "What must happen before merger completes"},
                {"key": "representations", "label": "Representations and Warranties", "type": "textarea", "placeholder": "Financial accuracy, legal compliance, no material adverse changes", "help": "Promises each company makes"},
                {"key": "board_composition", "label": "Post-Merger Board Composition", "type": "textarea", "placeholder": "5 members from ABC, 3 from XYZ", "help": "How board will be structured"},
                {"key": "employee_treatment", "label": "Employee Treatment", "type": "textarea", "placeholder": "All employees retained for 12 months, benefits maintained", "help": "How employees will be handled"},
                {"key": "termination", "label": "Termination Rights", "type": "textarea", "placeholder": "Material breach, failure to obtain approvals", "help": "When parties can walk away"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of Delaware", "help": "Applicable jurisdiction"},
            ],
            "Franchise Agreement": [
                {"key": "franchisor", "label": "Franchisor Name", "type": "text", "placeholder": "FastFood Brands Inc.", "help": "Company granting the franchise"},
                {"key": "franchisee", "label": "Franchisee Name", "type": "text", "placeholder": "John Doe", "help": "Individual/entity receiving franchise rights"},
                {"key": "territory", "label": "Territory", "type": "text", "placeholder": "Downtown Seattle, Washington", "help": "Geographic area of franchise"},
                {"key": "franchise_fee", "label": "Initial Franchise Fee", "type": "text", "placeholder": "$50,000", "help": "Upfront fee to acquire franchise"},
                {"key": "royalty_rate", "label": "Royalty Rate", "type": "text", "placeholder": "6% of gross sales", "help": "Ongoing royalty payments"},
                {"key": "marketing_fee", "label": "Marketing Fee", "type": "text", "placeholder": "2% of gross sales", "help": "Contribution to marketing fund"},
                {"key": "term", "label": "Initial Term", "type": "text", "placeholder": "10 years", "help": "Length of franchise agreement"},
                {"key": "renewal", "label": "Renewal Terms", "type": "text", "placeholder": "Two 5-year renewal options", "help": "Ability to extend franchise"},
                {"key": "training", "label": "Training Requirements", "type": "textarea", "placeholder": "2 weeks initial training at headquarters, ongoing support", "help": "Training franchisor will provide"},
                {"key": "standards", "label": "Operating Standards", "type": "textarea", "placeholder": "Must follow operations manual, quality standards, brand guidelines", "help": "Requirements franchisee must meet"},
                {"key": "termination", "label": "Termination Conditions", "type": "textarea", "placeholder": "Material breach, bankruptcy, failure to meet standards", "help": "When franchise can be terminated"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of Illinois", "help": "Applicable jurisdiction"},
            ],
            "Consulting Agreement": [
                {"key": "client", "label": "Client Name", "type": "text", "placeholder": "ABC Corporation", "help": "Company receiving consulting services"},
                {"key": "consultant", "label": "Consultant Name", "type": "text", "placeholder": "Jane Doe Consulting LLC", "help": "Individual or firm providing services"},
                {"key": "services", "label": "Services Description", "type": "textarea", "placeholder": "Strategic planning, market analysis, implementation support", "help": "Detailed scope of consulting work"},
                {"key": "deliverables", "label": "Deliverables", "type": "textarea", "placeholder": "Monthly reports, strategic plan document, presentation to board", "help": "Specific outputs consultant will provide"},
                {"key": "fee_structure", "label": "Fee Structure", "type": "text", "placeholder": "Hourly rate, fixed fee, or retainer", "help": "How consultant will be paid"},
                {"key": "compensation", "label": "Compensation Amount", "type": "text", "placeholder": "$200/hour, $50,000 project fee", "help": "Payment amount"},
                {"key": "payment_terms", "label": "Payment Terms", "type": "text", "placeholder": "Net 15 days from invoice", "help": "When payment is due"},
                {"key": "expenses", "label": "Expense Reimbursement", "type": "text", "placeholder": "Reasonable pre-approved expenses", "help": "How expenses are handled"},
                {"key": "start_date", "label": "Start Date", "type": "date", "help": "When consulting engagement begins"},
                {"key": "end_date", "label": "End Date", "type": "date", "help": "When engagement is expected to end"},
                {"key": "confidentiality", "label": "Confidentiality Terms", "type": "textarea", "placeholder": "Consultant shall maintain strict confidentiality", "help": "Protection of client information"},
                {"key": "ip_ownership", "label": "IP Ownership", "type": "text", "placeholder": "Client owns all work product", "help": "Who owns deliverables"},
                {"key": "termination", "label": "Termination Notice", "type": "text", "placeholder": "30 days written notice", "help": "How agreement can be ended"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of California", "help": "Applicable jurisdiction"},
            ],
            "Service Level Agreement (SLA)": [
                {"key": "provider", "label": "Service Provider", "type": "text", "placeholder": "CloudTech Services Inc.", "help": "Company providing the services"},
                {"key": "client", "label": "Client", "type": "text", "placeholder": "ABC Corporation", "help": "Company receiving services"},
                {"key": "services", "label": "Services Covered", "type": "textarea", "placeholder": "Cloud hosting, database management, technical support", "help": "Which services this SLA covers"},
                {"key": "uptime_guarantee", "label": "Uptime Guarantee", "type": "text", "placeholder": "99.9% uptime", "help": "Guaranteed availability percentage"},
                {"key": "response_time", "label": "Response Time", "type": "text", "placeholder": "Critical: 1 hour, High: 4 hours, Normal: 24 hours", "help": "How quickly provider responds to issues"},
                {"key": "resolution_time", "label": "Resolution Time", "type": "text", "placeholder": "Critical: 4 hours, High: 1 business day, Normal: 3 business days", "help": "Target time to fix issues"},
                {"key": "support_hours", "label": "Support Hours", "type": "text", "placeholder": "24/7/365 for critical, business hours for normal", "help": "When support is available"},
                {"key": "monitoring", "label": "Monitoring and Reporting", "type": "textarea", "placeholder": "Real-time monitoring, monthly performance reports", "help": "How performance is tracked"},
                {"key": "penalties", "label": "Service Credits/Penalties", "type": "textarea", "placeholder": "5% credit for each 0.1% below 99.9% uptime", "help": "Remedies if SLA not met"},
                {"key": "escalation", "label": "Escalation Procedures", "type": "textarea", "placeholder": "Level 1: Support team, Level 2: Manager, Level 3: VP Operations", "help": "How issues are escalated"},
                {"key": "maintenance_windows", "label": "Maintenance Windows", "type": "text", "placeholder": "Sunday 2-6 AM EST", "help": "Scheduled downtime periods"},
                {"key": "term", "label": "SLA Term", "type": "text", "placeholder": "12 months, auto-renewing", "help": "How long SLA is in effect"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of Virginia", "help": "Applicable jurisdiction"},
            ],
            "Board Resolution": [
                {"key": "company_name", "label": "Company Name", "type": "text", "placeholder": "ABC Corporation", "help": "Name of the corporation"},
                {"key": "meeting_date", "label": "Meeting Date", "type": "date", "help": "Date of board meeting"},
                {"key": "meeting_type", "label": "Meeting Type", "type": "text", "placeholder": "Regular / Special / Written Consent", "help": "Type of board meeting"},
                {"key": "attendees", "label": "Directors Present", "type": "textarea", "placeholder": "John Smith (Chair), Jane Doe, Bob Johnson", "help": "Board members attending"},
                {"key": "quorum", "label": "Quorum Established", "type": "text", "placeholder": "Yes / No", "help": "Was quorum present?"},
                {"key": "resolution_title", "label": "Resolution Title", "type": "text", "placeholder": "Approval of Annual Budget", "help": "Subject of the resolution"},
                {"key": "whereas_clauses", "label": "WHEREAS Clauses", "type": "textarea", "placeholder": "Background and reasons for the resolution", "help": "Recitals explaining context"},
                {"key": "resolved_clauses", "label": "RESOLVED Clauses", "type": "textarea", "placeholder": "Specific actions being authorized", "help": "The actual decisions/actions"},
                {"key": "vote_result", "label": "Vote Result", "type": "text", "placeholder": "Unanimous / 3 in favor, 0 opposed, 0 abstaining", "help": "Outcome of the vote"},
                {"key": "authority_granted", "label": "Authority Granted To", "type": "text", "placeholder": "CEO, CFO, Secretary", "help": "Who is authorized to act"},
                {"key": "effective_date", "label": "Effective Date", "type": "date", "help": "When resolution takes effect"},
                {"key": "secretary_certification", "label": "Secretary Certification", "type": "textarea", "placeholder": "I hereby certify that the foregoing is a true copy", "help": "Secretary's attestation"},
            ],
            # ===== EMPLOYMENT & HR =====
            "Severance Agreement": [
                {"key": "employer", "label": "Employer Name", "type": "text", "placeholder": "ABC Corporation", "help": "Company providing severance"},
                {"key": "employee", "label": "Employee Name", "type": "text", "placeholder": "John Doe", "help": "Employee receiving severance"},
                {"key": "position", "label": "Employee Position", "type": "text", "placeholder": "Senior Vice President", "help": "Employee's job title"},
                {"key": "termination_date", "label": "Termination Date", "type": "date", "help": "Last day of employment"},
                {"key": "severance_pay", "label": "Severance Payment", "type": "text", "placeholder": "12 months base salary ($180,000)", "help": "Cash severance amount"},
                {"key": "payment_schedule", "label": "Payment Schedule", "type": "text", "placeholder": "Lump sum or monthly installments", "help": "How severance is paid"},
                {"key": "benefits_continuation", "label": "Benefits Continuation", "type": "textarea", "placeholder": "Health insurance for 12 months, vested stock options exercisable for 90 days", "help": "Ongoing benefits"},
                {"key": "bonus_treatment", "label": "Bonus Treatment", "type": "text", "placeholder": "Pro-rated annual bonus for time worked", "help": "How bonuses are handled"},
                {"key": "outplacement", "label": "Outplacement Services", "type": "text", "placeholder": "6 months executive outplacement", "help": "Career transition assistance"},
                {"key": "release_of_claims", "label": "Release of Claims", "type": "textarea", "placeholder": "Employee releases all claims against employer", "help": "Legal claims employee gives up"},
                {"key": "non_disparagement", "label": "Non-Disparagement", "type": "text", "placeholder": "Mutual agreement not to make negative statements", "help": "Restrictions on public comments"},
                {"key": "confidentiality", "label": "Confidentiality", "type": "text", "placeholder": "Terms of agreement remain confidential", "help": "Non-disclosure of severance terms"},
                {"key": "cooperation", "label": "Post-Termination Cooperation", "type": "textarea", "placeholder": "Reasonable cooperation with transition, litigation", "help": "Ongoing obligations"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of New York", "help": "Applicable jurisdiction"},
            ],
            "Offer Letter": [
                {"key": "company", "label": "Company Name", "type": "text", "placeholder": "ABC Corporation", "help": "Employer offering the position"},
                {"key": "candidate", "label": "Candidate Name", "type": "text", "placeholder": "Jane Smith", "help": "Person receiving the offer"},
                {"key": "position", "label": "Position Title", "type": "text", "placeholder": "Senior Software Engineer", "help": "Job title being offered"},
                {"key": "department", "label": "Department", "type": "text", "placeholder": "Engineering", "help": "Which team/department"},
                {"key": "reports_to", "label": "Reports To", "type": "text", "placeholder": "VP of Engineering", "help": "Direct supervisor"},
                {"key": "start_date", "label": "Start Date", "type": "date", "help": "Expected first day of work"},
                {"key": "salary", "label": "Annual Salary", "type": "text", "placeholder": "$150,000", "help": "Base annual compensation"},
                {"key": "bonus", "label": "Bonus Opportunity", "type": "text", "placeholder": "15% annual target bonus", "help": "Variable compensation"},
                {"key": "equity", "label": "Equity Grant", "type": "text", "placeholder": "50,000 stock options, 4-year vesting", "help": "Stock options or RSUs"},
                {"key": "benefits", "label": "Benefits Summary", "type": "textarea", "placeholder": "Health, dental, vision, 401k matching, PTO", "help": "Employee benefits package"},
                {"key": "work_location", "label": "Work Location", "type": "text", "placeholder": "San Francisco, CA (Hybrid)", "help": "Where work will be performed"},
                {"key": "employment_type", "label": "Employment Type", "type": "text", "placeholder": "Full-time, At-Will", "help": "Employment classification"},
                {"key": "contingencies", "label": "Offer Contingencies", "type": "textarea", "placeholder": "Background check, reference verification, I-9 verification", "help": "Conditions that must be met"},
                {"key": "expiration", "label": "Offer Expiration", "type": "date", "help": "Deadline to accept offer"},
            ],
            "Non-Compete Agreement": [
                {"key": "employer", "label": "Employer Name", "type": "text", "placeholder": "ABC Corporation", "help": "Company being protected"},
                {"key": "employee", "label": "Employee Name", "type": "text", "placeholder": "John Doe", "help": "Person agreeing not to compete"},
                {"key": "position", "label": "Position", "type": "text", "placeholder": "Sales Director", "help": "Employee's role"},
                {"key": "effective_date", "label": "Effective Date", "type": "date", "help": "When agreement begins"},
                {"key": "restricted_activities", "label": "Restricted Activities", "type": "textarea", "placeholder": "Employment by competitor, soliciting clients, competing business", "help": "What employee cannot do"},
                {"key": "duration", "label": "Restriction Duration", "type": "text", "placeholder": "24 months after employment ends", "help": "How long restrictions last"},
                {"key": "geographic_scope", "label": "Geographic Scope", "type": "text", "placeholder": "50-mile radius of employer's offices", "help": "Where restrictions apply"},
                {"key": "industry_scope", "label": "Industry Scope", "type": "textarea", "placeholder": "Software development for healthcare industry", "help": "Which industries are restricted"},
                {"key": "consideration", "label": "Consideration", "type": "text", "placeholder": "Continued employment, $10,000 signing bonus", "help": "What employee receives in exchange"},
                {"key": "non_solicitation", "label": "Non-Solicitation Terms", "type": "textarea", "placeholder": "May not solicit employees or clients for 18 months", "help": "Restrictions on soliciting"},
                {"key": "remedies", "label": "Remedies for Breach", "type": "textarea", "placeholder": "Injunctive relief, monetary damages, attorney fees", "help": "What happens if violated"},
                {"key": "severability", "label": "Severability Clause", "type": "text", "placeholder": "Court may modify overly broad restrictions", "help": "If parts are unenforceable"},
                {"key": "governing_law", "label": "Governing Law", "type": "text", "placeholder": "State of Texas", "help": "Applicable jurisdiction"},
            ],
            "Employee Handbook": [
                {"key": "company_name", "label": "Company Name", "type": "text", "placeholder": "ABC Corporation", "help": "Name of the organization"},
                {"key": "effective_date", "label": "Effective Date", "type": "date", "help": "When handbook takes effect"},
                {"key": "mission_values", "label": "Mission and Values", "type": "textarea", "placeholder": "Company mission statement and core values", "help": "Company culture and principles"},
                {"key": "employment_policies", "label": "Employment Policies", "type": "textarea", "placeholder": "At-will employment, equal opportunity, anti-discrimination", "help": "Core employment policies"},
                {"key": "work_schedule", "label": "Work Schedule and Hours", "type": "textarea", "placeholder": "Standard hours, overtime, flex time, remote work", "help": "When and where employees work"},
                {"key": "compensation", "label": "Compensation Policies", "type": "textarea", "placeholder": "Pay periods, overtime, bonuses, raises", "help": "How employees are paid"},
                {"key": "benefits", "label": "Benefits Overview", "type": "textarea", "placeholder": "Health insurance, retirement, PTO, holidays", "help": "Employee benefits available"},
                {"key": "pto_policy", "label": "PTO/Leave Policy", "type": "textarea", "placeholder": "Vacation, sick leave, parental leave, FMLA", "help": "Time off policies"},
                {"key": "conduct_standards", "label": "Standards of Conduct", "type": "textarea", "placeholder": "Expected behavior, dress code, attendance", "help": "Behavioral expectations"},
                {"key": "technology_use", "label": "Technology and Social Media", "type": "textarea", "placeholder": "Acceptable use of company tech, social media guidelines", "help": "Technology policies"},
                {"key": "safety", "label": "Health and Safety", "type": "textarea", "placeholder": "Workplace safety, reporting injuries, emergency procedures", "help": "Safety protocols"},
                {"key": "harassment_policy", "label": "Anti-Harassment Policy", "type": "textarea", "placeholder": "Zero tolerance, reporting procedures, investigation process", "help": "Harassment prevention"},
                {"key": "discipline", "label": "Disciplinary Procedures", "type": "textarea", "placeholder": "Progressive discipline, grounds for termination", "help": "How violations are handled"},
                {"key": "acknowledgment", "label": "Acknowledgment Language", "type": "textarea", "placeholder": "Employee acknowledgment that they received and read handbook", "help": "Receipt confirmation"},
            ],
        }

        # Merge all document types
        ALL_EXTENDED_TYPES = {**EXTENDED_DOCUMENT_TYPES, **ADDITIONAL_DOCUMENT_TYPES}

        # Check extended types first
        if document_type in ALL_EXTENDED_TYPES:
            return ALL_EXTENDED_TYPES[document_type]

        # Check base DOCUMENT_TYPES
        if document_type in DOCUMENT_TYPES:
            return DOCUMENT_TYPES[document_type]["fields"]

        # For truly unknown types, ask the LLM
        prompt = f"""What fields/information are typically required to draft a {document_type}?

Return a JSON array of field objects, each with:
- "key": snake_case field identifier
- "label": Human-readable field label
- "type": "text" or "date" or "textarea"
- "placeholder": Example value
- "help": Brief explanation of what to enter

Example:
[
    {{"key": "party_a", "label": "First Party Name", "type": "text", "placeholder": "John Doe", "help": "Full legal name"}},
    {{"key": "effective_date", "label": "Effective Date", "type": "date", "placeholder": "", "help": "When the agreement takes effect"}}
]

Return ONLY the JSON array, no other text."""

        messages = [
            {"role": "system", "content": "You are a legal document expert. Return only valid JSON."},
            {"role": "user", "content": prompt}
        ]

        try:
            response = self.llm.chat(messages, temperature=0.2, max_tokens=2000)
            import json
            # Extract JSON array
            json_start = response.find('[')
            json_end = response.rfind(']') + 1
            if json_start >= 0 and json_end > json_start:
                fields = json.loads(response[json_start:json_end])
                return fields
            else:
                # Fallback
                return EXTENDED_DOCUMENT_TYPES["Custom Document"]
        except Exception:
            return EXTENDED_DOCUMENT_TYPES["Custom Document"]

    # -- Style examples from ChromaDB ------------------------------------

    def get_style_examples(
        self, document_type: str, n_examples: int = 3
    ) -> list[str]:
        """
        Retrieve example chunks from ChromaDB filtered by document_type.
        Prioritizes chunks from indexed real documents over manually uploaded examples.
        """
        if self.collection is None or self.collection.count() == 0:
            return []
        try:
            # Try to get examples from indexed documents (from scanner)
            results = self.collection.query(
                query_texts=[f"{document_type} professional legal document"],
                n_results=min(n_examples * 2, self.collection.count()),  # Get more candidates
                where={"document_type": document_type},
            )

            if results["documents"] and results["documents"][0]:
                # Filter for higher-quality examples (longer chunks, from indexed files)
                candidates = results["documents"][0]
                metadatas = results.get("metadatas", [[]])[0]

                # Prioritize chunks from scanned documents (have source_path metadata)
                scored_examples = []
                for chunk, metadata in zip(candidates, metadatas):
                    score = len(chunk)  # Base score on length
                    if metadata.get("source_path"):
                        score += 500  # Bonus for being from scanned document
                    if metadata.get("confidence", 0) > 0.7:
                        score += 200  # Bonus for high classification confidence
                    scored_examples.append((score, chunk))

                # Sort by score and take top n
                scored_examples.sort(reverse=True, key=lambda x: x[0])
                return [chunk for _, chunk in scored_examples[:n_examples]]

            return []
        except Exception:
            # If where-filter fails (no matching docs), fall back to empty
            return []

    # -- Reference data --------------------------------------------------

    def fetch_reference_data(
        self,
        document_type: str,
        params: dict,
        sec_client=None,
    ) -> str:
        """Optionally pull reference data from SEC EDGAR."""
        parts: list[str] = []

        # SEC EDGAR — useful for filings & contracts mentioning public companies
        if sec_client and sec_client.is_configured():
            try:
                cik = params.get("sec_cik", "").strip()
                if cik:
                    filings = sec_client.get_company_filings(cik)
                    name = filings.get("name", "")
                    recent = filings.get("filings", {}).get("recent", {})
                    forms = recent.get("form", [])[:5]
                    dates = recent.get("filingDate", [])[:5]
                    lines = [f"Company: {name}"]
                    for f, d in zip(forms, dates):
                        lines.append(f"  - {f} filed {d}")
                    parts.append(
                        "SEC EDGAR reference data:\n" + "\n".join(lines)
                    )
                else:
                    # Try a keyword search based on entity name
                    entity = params.get("entity_name", "") or params.get("party_a", "")
                    if entity:
                        hits = sec_client.search_filings(entity, max_results=3)
                        if hits:
                            lines = [
                                f"  - {h['entity_name']}: {h['form_type']} ({h['file_date']})"
                                for h in hits
                            ]
                            parts.append(
                                "SEC EDGAR search results:\n" + "\n".join(lines)
                            )
            except Exception as e:
                parts.append(f"(SEC EDGAR lookup failed: {e})")

        return "\n\n".join(parts)

    # -- Prompt construction ---------------------------------------------

    def build_prompt(
        self,
        document_type: str,
        params: dict,
        style_examples: list[str],
        reference_data: str,
    ) -> tuple[str, str]:
        """Return (system_prompt, user_prompt) for the LLM."""
        doc_def = DOCUMENT_TYPES[document_type]
        system_prompt = doc_def["system_prompt"]

        # -- Build user prompt -------------------------------------------
        sections: list[str] = []

        # Style examples
        if style_examples:
            examples_text = "\n---\n".join(style_examples[:3])
            sections.append(
                "STYLE REFERENCE — follow the tone, structure, and formatting "
                "of these example excerpts:\n\n" + examples_text
            )

        # Reference data
        if reference_data:
            sections.append("REFERENCE DATA:\n" + reference_data)

        # User parameters
        param_lines: list[str] = []
        for field in doc_def["fields"]:
            key = field["key"]
            value = params.get(key, "")
            if value:
                # Format dates nicely
                if hasattr(value, "strftime"):
                    value = value.strftime("%B %d, %Y")
                param_lines.append(f"{field['label']}: {value}")

        sections.append(
            f"Generate a {doc_def['label']} with the following parameters:\n\n"
            + "\n".join(param_lines)
        )

        user_prompt = "\n\n".join(sections)
        return system_prompt, user_prompt

    # -- Full pipeline ---------------------------------------------------

    def generate(
        self,
        document_type: str,
        params: dict,
        sec_client=None,
        use_sec: bool = False,
    ) -> str:
        """Run the full generation pipeline and return the document text."""
        # 1. Style examples
        style_examples = self.get_style_examples(document_type)

        # 2. Reference data
        reference_data = ""
        if use_sec:
            reference_data = self.fetch_reference_data(
                document_type,
                params,
                sec_client=sec_client if use_sec else None,
            )

        # 3. Build prompt
        system_prompt, user_prompt = self.build_prompt(
            document_type, params, style_examples, reference_data
        )

        # 4. Generate
        return self.llm.generate_document(system_prompt, user_prompt)

    # -- DOCX conversion -------------------------------------------------

    @staticmethod
    def text_to_docx(text: str, title: str = "Generated Document") -> bytes:
        """Convert plain text to a formatted .docx and return the bytes."""
        doc = DocxDocument()

        # Page setup
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run(
            "DRAFT — FOR ATTORNEY REVIEW ONLY — NOT LEGAL ADVICE"
        )
        run.bold = True
        run.font.size = Pt(9)

        doc.add_paragraph("")  # spacer

        # Body text — preserve paragraphs and basic structure
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # Detect heading-like lines (ALL CAPS or starts with a number + period)
            if (
                stripped.isupper() and len(stripped) < 120
            ) or re.match(r"^(ARTICLE|SECTION|RECITAL)\s", stripped, re.IGNORECASE):
                doc.add_heading(stripped, level=1)
            elif re.match(r"^\d+\.\s+[A-Z]", stripped):
                doc.add_heading(stripped, level=2)
            else:
                p = doc.add_paragraph(stripped)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    run.font.size = Pt(11)

        # Write to bytes
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def replace_text_preserve_word_template(original_file_bytes: bytes, new_text: str) -> bytes:
        """
        Replace text inside an existing Word package while preserving package assets.

        This keeps original styles, headers/footers, and docm macros in the package.
        Text replacement is best-effort and maps new text across existing text runs.
        """
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        source = io.BytesIO(original_file_bytes)
        out_buf = io.BytesIO()

        with zipfile.ZipFile(source, "r") as zin:
            if "word/document.xml" not in zin.namelist():
                raise ValueError("Invalid Word package: missing word/document.xml")

            xml_bytes = zin.read("word/document.xml")
            root = ET.fromstring(xml_bytes)
            text_nodes = root.findall(".//w:t", ns)
            if not text_nodes:
                raise ValueError("No editable text nodes found in document.xml")

            flat_text = new_text.replace("\r\n", "\n").replace("\r", "\n").replace("\n", " ")
            original_lengths = [len(node.text or "") for node in text_nodes]
            cursor = 0

            for node, chunk_len in zip(text_nodes, original_lengths):
                if chunk_len <= 0:
                    node.text = ""
                    continue
                node.text = flat_text[cursor:cursor + chunk_len]
                cursor += chunk_len

            if cursor < len(flat_text):
                text_nodes[-1].text = (text_nodes[-1].text or "") + flat_text[cursor:]

            ET.register_namespace("w", ns["w"])
            updated_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

            with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == "word/document.xml":
                        zout.writestr(item, updated_xml)
                    else:
                        zout.writestr(item, zin.read(item.filename))

        return out_buf.getvalue()

    @staticmethod
    def diff_to_docx(original_text: str, revised_text: str, title: str = "Document Redline") -> bytes:
        """Create a native Word tracked-changes .docx with insertions/deletions."""

        def _tokens(content: str) -> list[str]:
            return re.findall(r"\s+|[A-Za-z0-9_]+|[^\w\s]", content or "")

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns_xml = "http://www.w3.org/XML/1998/namespace"
        qn = lambda tag: f"{{{ns_w}}}{tag}"
        now_iso = date.today().isoformat() + "T00:00:00Z"

        def _set_space_preserve(el: ET.Element, value: str):
            if value and (value != value.strip() or "  " in value):
                el.set(f"{{{ns_xml}}}space", "preserve")

        def _append_text_container(parent: ET.Element, content: str, *, deleted: bool):
            if content == "":
                return
            r = ET.SubElement(parent, qn("r"))
            t = ET.SubElement(r, qn("delText" if deleted else "t"))
            _set_space_preserve(t, content)
            t.text = content

        def _append_break(parent: ET.Element):
            r = ET.SubElement(parent, qn("r"))
            ET.SubElement(r, qn("br"))

        def _append_segment(parent_p: ET.Element, segment: str, mode: str, rev_id: int):
            parts = segment.split("\n")

            if mode == "equal":
                for idx, part in enumerate(parts):
                    _append_text_container(parent_p, part, deleted=False)
                    if idx < len(parts) - 1:
                        _append_break(parent_p)
                return rev_id

            rev_tag = "ins" if mode == "insert" else "del"
            rev_el = ET.SubElement(
                parent_p,
                qn(rev_tag),
                {
                    qn("id"): str(rev_id),
                    qn("author"): "Document KB AI",
                    qn("date"): now_iso,
                },
            )
            for idx, part in enumerate(parts):
                _append_text_container(rev_el, part, deleted=(mode == "delete"))
                if idx < len(parts) - 1:
                    _append_break(rev_el)
            return rev_id + 1

        base_doc = DocxDocument()
        heading = base_doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        base_doc.add_paragraph("Open in Microsoft Word and use Review > Accept/Reject to finalize changes.")
        body_para = base_doc.add_paragraph("")

        base_buf = io.BytesIO()
        base_doc.save(base_buf)

        source = io.BytesIO(base_buf.getvalue())
        out_buf = io.BytesIO()

        with zipfile.ZipFile(source, "r") as zin:
            if "word/document.xml" not in zin.namelist():
                raise ValueError("Invalid Word package: missing word/document.xml")

            doc_root = ET.fromstring(zin.read("word/document.xml"))
            paragraphs = doc_root.findall(f".//{qn('p')}")
            if not paragraphs:
                raise ValueError("Word document has no paragraph container")

            target_p = paragraphs[-1]
            for child in list(target_p):
                target_p.remove(child)

            sm = difflib.SequenceMatcher(a=_tokens(original_text), b=_tokens(revised_text))
            rev_id = 1
            for op, i1, i2, j1, j2 in sm.get_opcodes():
                if op == "equal":
                    seg = "".join(sm.b[j1:j2])
                    rev_id = _append_segment(target_p, seg, "equal", rev_id)
                elif op == "insert":
                    seg = "".join(sm.b[j1:j2])
                    rev_id = _append_segment(target_p, seg, "insert", rev_id)
                elif op == "delete":
                    seg = "".join(sm.a[i1:i2])
                    rev_id = _append_segment(target_p, seg, "delete", rev_id)
                elif op == "replace":
                    del_seg = "".join(sm.a[i1:i2])
                    ins_seg = "".join(sm.b[j1:j2])
                    rev_id = _append_segment(target_p, del_seg, "delete", rev_id)
                    rev_id = _append_segment(target_p, ins_seg, "insert", rev_id)

            ET.register_namespace("w", ns_w)
            updated_document_xml = ET.tostring(doc_root, encoding="utf-8", xml_declaration=True)

            updated_settings_xml = None
            if "word/settings.xml" in zin.namelist():
                settings_root = ET.fromstring(zin.read("word/settings.xml"))
                has_track = settings_root.find(f".//{qn('trackRevisions')}") is not None
                if not has_track:
                    ET.SubElement(settings_root, qn("trackRevisions"))
                ET.register_namespace("w", ns_w)
                updated_settings_xml = ET.tostring(settings_root, encoding="utf-8", xml_declaration=True)

            with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == "word/document.xml":
                        zout.writestr(item, updated_document_xml)
                    elif item.filename == "word/settings.xml" and updated_settings_xml is not None:
                        zout.writestr(item, updated_settings_xml)
                    else:
                        zout.writestr(item, zin.read(item.filename))

        return out_buf.getvalue()
