import os
from typing import List, Dict, Any
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from config import Config

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
        )

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error extracting text from TXT: {str(e)}")

    def extract_text(self, file_path: str) -> str:
        """Extract text from file based on extension"""
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """Process a document and return chunks with metadata"""
        # Extract text
        text = self.extract_text(file_path)

        if not text.strip():
            raise ValueError(f"No text could be extracted from {file_path}")

        # Create metadata
        filename = os.path.basename(file_path)
        metadata = {
            "source": filename,
            "file_path": file_path,
            "file_type": os.path.splitext(filename)[1].lower()
        }

        # Split text into chunks
        chunks = self.text_splitter.split_text(text)

        # Create LangChain documents with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata["chunk_id"] = i
            chunk_metadata["total_chunks"] = len(chunks)

            documents.append(LangchainDocument(
                page_content=chunk,
                metadata=chunk_metadata
            ))

        return documents

    def process_multiple_documents(self, file_paths: List[str]) -> List[LangchainDocument]:
        """Process multiple documents and return all chunks"""
        all_documents = []

        for file_path in file_paths:
            try:
                documents = self.process_document(file_path)
                all_documents.extend(documents)
                print(f"Processed {file_path}: {len(documents)} chunks")
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
                continue

        return all_documents