"""
Document Export Module
Exports generated documents to various formats (DOCX, PDF, TXT)
"""

import os
import tempfile
from typing import Optional, Dict, Any
from datetime import datetime
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class DocumentExporter:
    """Export documents to various formats"""

    def __init__(self):
        self.temp_dir = tempfile.gettempdir()

    def export_to_txt(self, content: str, filename: Optional[str] = None) -> str:
        """Export document to plain text format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.txt"

        file_path = os.path.join(self.temp_dir, filename)

        # Clean up the content for plain text
        clean_content = self._clean_text_for_export(content)

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(clean_content)

        return file_path

    def export_to_docx(self, content: str, filename: Optional[str] = None,
                      document_type: str = "general") -> Optional[str]:
        """Export document to Microsoft Word format"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.docx"

        file_path = os.path.join(self.temp_dir, filename)

        try:
            # Create Word document
            doc = Document()

            # Set up styles based on document type
            self._setup_docx_styles(doc, document_type)

            # Parse and add content
            self._add_content_to_docx(doc, content, document_type)

            # Save document
            doc.save(file_path)
            return file_path

        except Exception as e:
            print(f"Error creating DOCX: {e}")
            return None

    def export_to_pdf(self, content: str, filename: Optional[str] = None,
                     document_type: str = "general") -> Optional[str]:
        """Export document to PDF format"""
        if not PDF_AVAILABLE:
            raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.pdf"

        file_path = os.path.join(self.temp_dir, filename)

        try:
            # Create PDF document
            doc = SimpleDocTemplate(file_path, pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)

            # Get styles
            styles = self._get_pdf_styles(document_type)

            # Parse content and create story
            story = self._create_pdf_story(content, styles)

            # Build PDF
            doc.build(story)
            return file_path

        except Exception as e:
            print(f"Error creating PDF: {e}")
            return None

    def _clean_text_for_export(self, content: str) -> str:
        """Clean text content for export"""
        # Remove any HTML-like tags if present
        content = re.sub(r'<[^>]+>', '', content)

        # Normalize whitespace
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        content = content.strip()

        # Add header with generation info
        header = f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n"
        header += "=" * 50 + "\n\n"

        return header + content

    def _setup_docx_styles(self, doc: Document, document_type: str):
        """Set up Word document styles based on document type"""
        styles = doc.styles

        # Create custom styles if they don't exist
        if document_type == "legal":
            self._create_legal_docx_styles(styles)
        elif document_type == "business":
            self._create_business_docx_styles(styles)
        else:
            self._create_general_docx_styles(styles)

    def _create_legal_docx_styles(self, styles):
        """Create legal document styles for Word"""
        try:
            # Legal heading style
            legal_heading = styles.add_style('Legal Heading', WD_STYLE_TYPE.PARAGRAPH)
            legal_heading.font.bold = True
            legal_heading.font.size = 14
            legal_heading.paragraph_format.space_after = 12
        except:
            pass  # Style might already exist

    def _create_business_docx_styles(self, styles):
        """Create business document styles for Word"""
        try:
            # Business heading style
            business_heading = styles.add_style('Business Heading', WD_STYLE_TYPE.PARAGRAPH)
            business_heading.font.bold = True
            business_heading.font.size = 12
            business_heading.paragraph_format.space_after = 6
        except:
            pass

    def _create_general_docx_styles(self, styles):
        """Create general document styles for Word"""
        # Use default styles for general documents
        pass

    def _add_content_to_docx(self, doc: Document, content: str, document_type: str):
        """Add content to Word document with appropriate formatting"""
        lines = content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                # Add spacing for empty lines
                doc.add_paragraph()
                continue

            # Detect headers (all caps, short lines, or lines ending with :)
            if self._is_header(line):
                if document_type == "legal":
                    p = doc.add_paragraph(line, style='Legal Heading')
                elif document_type == "business":
                    p = doc.add_paragraph(line, style='Business Heading')
                else:
                    p = doc.add_paragraph(line, style='Heading 2')
            # Detect bullet points
            elif re.match(r'^[•\-\*]\s+', line):
                p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            # Detect numbered lists
            elif re.match(r'^\d+[\.\)]\s+', line):
                p = doc.add_paragraph(line, style='List Number')
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)

                # Apply special formatting for legal documents
                if document_type == "legal":
                    # Look for legal terms to emphasize
                    if any(term in line.upper() for term in ['WHEREAS', 'THEREFORE', 'PARTY', 'PARTIES']):
                        p.runs[0].bold = True

    def _is_header(self, line: str) -> bool:
        """Determine if a line should be treated as a header"""
        return (
            len(line) < 100 and
            (line.isupper() or line.istitle() or line.endswith(':')) and
            not line.endswith('.')
        )

    def _get_pdf_styles(self, document_type: str) -> Dict[str, Any]:
        """Get PDF styles based on document type"""
        styles = getSampleStyleSheet()

        if document_type == "legal":
            # Create legal document styles
            legal_heading = ParagraphStyle(
                'LegalHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            styles.add(legal_heading)

        elif document_type == "business":
            # Create business document styles
            business_heading = ParagraphStyle(
                'BusinessHeading',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=6,
                fontName='Helvetica-Bold'
            )
            styles.add(business_heading)

        return styles

    def _create_pdf_story(self, content: str, styles) -> list:
        """Create PDF story from content"""
        story = []
        lines = content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 12))
                continue

            # Determine paragraph style
            if self._is_header(line):
                if 'LegalHeading' in styles:
                    story.append(Paragraph(line, styles['LegalHeading']))
                elif 'BusinessHeading' in styles:
                    story.append(Paragraph(line, styles['BusinessHeading']))
                else:
                    story.append(Paragraph(line, styles['Heading2']))
            elif re.match(r'^[•\-\*]\s+', line):
                # Bullet point
                story.append(Paragraph(f"• {line[2:].strip()}", styles['Normal']))
            else:
                story.append(Paragraph(line, styles['Normal']))

        return story

    def export_document(self, content: str, format_type: str = "txt",
                       filename: Optional[str] = None,
                       document_type: str = "general") -> Dict[str, Any]:
        """Export document to specified format"""

        try:
            if format_type.lower() == "txt":
                file_path = self.export_to_txt(content, filename)
                success = True
                error = None

            elif format_type.lower() == "docx":
                if not DOCX_AVAILABLE:
                    return {
                        "success": False,
                        "file_path": None,
                        "error": "DOCX export requires python-docx package. Install with: pip install python-docx"
                    }
                file_path = self.export_to_docx(content, filename, document_type)
                success = file_path is not None
                error = "Failed to create DOCX file" if not success else None

            elif format_type.lower() == "pdf":
                if not PDF_AVAILABLE:
                    return {
                        "success": False,
                        "file_path": None,
                        "error": "PDF export requires reportlab package. Install with: pip install reportlab"
                    }
                file_path = self.export_to_pdf(content, filename, document_type)
                success = file_path is not None
                error = "Failed to create PDF file" if not success else None

            else:
                return {
                    "success": False,
                    "file_path": None,
                    "error": f"Unsupported format: {format_type}. Supported formats: txt, docx, pdf"
                }

            result = {
                "success": success,
                "file_path": file_path,
                "format": format_type,
                "filename": Path(file_path).name if file_path else None,
                "error": error
            }

            if success and file_path:
                # Add file size information
                file_size = os.path.getsize(file_path)
                result["file_size"] = file_size
                result["file_size_mb"] = round(file_size / (1024 * 1024), 2)

            return result

        except Exception as e:
            return {
                "success": False,
                "file_path": None,
                "error": f"Export error: {str(e)}"
            }

    def get_supported_formats(self) -> Dict[str, Any]:
        """Get information about supported export formats"""
        formats = {
            "txt": {
                "available": True,
                "description": "Plain text format",
                "extension": ".txt"
            },
            "docx": {
                "available": DOCX_AVAILABLE,
                "description": "Microsoft Word format",
                "extension": ".docx",
                "requirements": "python-docx" if not DOCX_AVAILABLE else None
            },
            "pdf": {
                "available": PDF_AVAILABLE,
                "description": "Portable Document Format",
                "extension": ".pdf",
                "requirements": "reportlab" if not PDF_AVAILABLE else None
            }
        }

        return {
            "supported_formats": formats,
            "available_count": sum(1 for f in formats.values() if f["available"]),
            "total_count": len(formats)
        }