"""
Corporate Law Document Generator ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Streamlit Application

Features:
- User authentication (login/registration)
- Per-user data isolation
- Role-based access (admin/user)
- Modern, professional UI
- Three-tab interface: Chat Q&A, Generate Document, Settings
- Admin panel for user management
"""

import os
import io
import re
import sys
import json
import subprocess
import html as html_lib
import difflib
import hashlib
import tempfile
from pathlib import Path
from urllib.parse import urlparse
from dotenv import load_dotenv
load_dotenv()

import streamlit as st
from typing import List
from datetime import date, datetime
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _env_bool(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _env_csv(name: str, default: str) -> set[str]:
    value = os.getenv(name, default)
    return {
        item.strip().lower()
        for item in value.split(",")
        if item.strip()
    }

import chromadb
from chromadb.utils import embedding_functions

from llm_backend import LLMBackend
from document_generator import DocumentGenerator, DOCUMENT_TYPES
from api_clients import SECEdgarClient

# Continuous learning pipeline imports
from knowledge_db import KnowledgeDB
from document_scanner import DocumentScanner
from background_scanner import ScannerManager

# Authentication imports
from auth import AuthManager, init_session_state, require_auth, get_user_collection_name, logout
from auth_ui import (
    render_login_page,
    render_registration_page,
    render_verification_page,
    render_admin_panel,
    render_profile_settings,
    render_auth_sidebar
)
from styles import get_custom_css, render_header, render_footer


# ======================================================================
# Page config
# ======================================================================
st.set_page_config(
    page_title="Corporate Law Document Generator",
    page_icon="ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Apply custom CSS
st.markdown(get_custom_css(), unsafe_allow_html=True)


# ======================================================================
# Environment detection and smart defaults
# ======================================================================

def is_streamlit_cloud() -> bool:
    """Detect if running on Streamlit Cloud."""
    return os.getenv("STREAMLIT_SHARING_MODE") is not None or \
           os.getenv("STREAMLIT_RUNTIME_ENV") == "cloud"


FORCE_OLLAMA_FOR_ALL_USERS = os.getenv("FORCE_OLLAMA_FOR_ALL_USERS", "false").lower() == "true"
SHARED_KNOWLEDGE_SCOPE = os.getenv("SHARED_KNOWLEDGE_SCOPE", "true").lower() == "true"
SHARED_SCOPE_USER_ID = int(os.getenv("SHARED_SCOPE_USER_ID", "0"))
SHARED_COLLECTION_NAME = os.getenv("SHARED_COLLECTION_NAME", "documents_shared")

ENFORCE_WEB_SECURITY = _env_bool("ENFORCE_WEB_SECURITY", True)
REQUIRE_HTTPS = _env_bool("REQUIRE_HTTPS", True)
ALLOW_INSECURE_LOCALHOST = _env_bool("ALLOW_INSECURE_LOCALHOST", True)
STRICT_HOST_VALIDATION = _env_bool("STRICT_HOST_VALIDATION", True)
TRUSTED_HOSTS = _env_csv(
    "TRUSTED_HOSTS",
    "localhost,127.0.0.1,::1"
)


def _get_request_headers() -> dict[str, str]:
    try:
        context = getattr(st, "context", None)
        headers = getattr(context, "headers", None)
        if not headers:
            return {}
        return {str(k).lower(): str(v) for k, v in dict(headers).items()}
    except Exception:
        return {}


def _extract_header_value(headers: dict[str, str], *keys: str) -> str:
    for key in keys:
        if key in headers and headers[key]:
            return headers[key].split(",")[0].strip()
    return ""


def _normalize_host(host_value: str) -> str:
    host = host_value.strip().lower()
    if host.startswith("[") and "]" in host:
        host = host[1:host.index("]")]
    elif ":" in host:
        host = host.split(":", 1)[0]
    return host


def enforce_web_access_security() -> None:
    if not ENFORCE_WEB_SECURITY:
        return

    headers = _get_request_headers()
    if not headers:
        return

    host_raw = _extract_header_value(headers, "x-forwarded-host", "host")
    host = _normalize_host(host_raw)
    proto = _extract_header_value(headers, "x-forwarded-proto", "x-scheme", "x-forwarded-protocol").lower()

    if not host:
        return

    local_hosts = {"localhost", "127.0.0.1", "::1"}
    is_local = host in local_hosts

    # Never block localhost loopback hosts, even if TRUSTED_HOSTS is customized.
    if STRICT_HOST_VALIDATION and TRUSTED_HOSTS and not is_local and host not in TRUSTED_HOSTS:
        st.error("Access denied: untrusted host header.")
        st.stop()

    if REQUIRE_HTTPS and not (ALLOW_INSECURE_LOCALHOST and is_local):
        if proto and proto != "https":
            st.error("Secure HTTPS is required for web access.")
            st.stop()
        if not proto and not is_local:
            st.error("Secure HTTPS is required for web access.")
            st.stop()


def get_default_llm_provider() -> str:
    """Auto-detect best LLM provider based on environment.
    Defaults to Ollama for local deployments - all users share the host's Ollama instance.
    """
    if FORCE_OLLAMA_FOR_ALL_USERS:
        return "ollama"

    try:
        if "LLM_PROVIDER" in st.secrets:
            return st.secrets.get("LLM_PROVIDER", "ollama")
    except Exception:
        pass

    provider = os.getenv("LLM_PROVIDER", "").lower()
    if provider in ["ollama", "openai", "hf_local"]:
        return provider

    if is_streamlit_cloud():
        return "openai"

    return "ollama"


def get_config_value(key: str, default: str = "") -> str:
    """Get config from Streamlit secrets, then env vars, then default."""
    try:
        if key in st.secrets:
            value = st.secrets[key]
            if value and value not in ["your-api-key-here", "your_openai_api_key_here"]:
                return value
    except Exception:
        pass
    return os.getenv(key, default)


# ======================================================================
# Initialize authentication
# ======================================================================

init_session_state(st.session_state)
enforce_web_access_security()
auth_manager = AuthManager()


# ======================================================================
# Session state defaults with smart environment detection
# ======================================================================

DEFAULT_PROVIDER = get_default_llm_provider()

_DEFAULTS = {
    "messages": [],
    "processed_files": set(),
    "llm_provider": DEFAULT_PROVIDER,
    "llm_model": (
        get_config_value("OPENAI_MODEL", "gpt-4o-mini")
        if DEFAULT_PROVIDER == "openai"
        else get_config_value("OLLAMA_MODEL", "llama3.1:8b")
    ),
    "llm_base_url": get_config_value("OLLAMA_BASE_URL", "http://localhost:11434/v1"),
    "openai_api_key": get_config_value("OPENAI_API_KEY", ""),
    "hf_local_model_path": get_config_value("HF_LOCAL_MODEL_PATH", ""),
    "hf_local_max_new_tokens": get_config_value("HF_LOCAL_MAX_NEW_TOKENS", "2048"),
    "sec_user_agent": get_config_value("SEC_EDGAR_USER_AGENT", ""),
    "generated_text": "",
    "generated_title": "",
    "onboarding_complete": False,
    "workflow_mode": None,  # None, "create", "edit", or "learn"
    "show_settings": False,
    "show_knowledge_base": False,
    "show_model_improvement": False,
}

for key, val in _DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = val

if FORCE_OLLAMA_FOR_ALL_USERS:
    st.session_state.llm_provider = "ollama"


def get_data_scope_user_id() -> int:
    """Return shared or per-user scope id for knowledge/indexing operations."""
    if SHARED_KNOWLEDGE_SCOPE:
        return SHARED_SCOPE_USER_ID
    if require_auth(st.session_state):
        return st.session_state.current_user.user_id
    return 0


# ======================================================================
# ChromaDB setup with per-user collections
# ======================================================================

@st.cache_resource
def get_chroma_client():
    return chromadb.PersistentClient(path="./chroma_db")


@st.cache_resource
def get_embedding_function():
    return embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name="all-MiniLM-L6-v2"
    )


@st.cache_resource
def get_knowledge_db():
    """Get the knowledge database instance."""
    return KnowledgeDB(db_path="./knowledge.db")


@st.cache_resource
def get_background_scanner(user_id: int, _llm, _collection, _knowledge_db):
    """Initialize and get the background scanner for the current user (manual/admin only)."""
    return ScannerManager.initialize(
        user_id=user_id,
        llm=_llm,
        chroma_collection=_collection,
        knowledge_db=_knowledge_db,
        auto_start=False,
    )


def get_collection():
    """Get the ChromaDB collection for the current scope (shared or per-user)."""
    if SHARED_KNOWLEDGE_SCOPE:
        collection_name = SHARED_COLLECTION_NAME
    elif not require_auth(st.session_state):
        collection_name = "documents"
    else:
        user = st.session_state.current_user
        collection_name = get_user_collection_name(user.user_id)

    client = get_chroma_client()
    ef = get_embedding_function()
    return client.get_or_create_collection(name=collection_name, embedding_function=ef)


# ======================================================================
# Helpers ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â LLM / clients
# ======================================================================

class _UnavailableLLM:
    def __init__(self, reason: str):
        self.reason = reason

    def is_available(self) -> bool:
        return False

    def chat(self, *args, **kwargs):
        raise RuntimeError(self.reason)

    def generate_document(self, *args, **kwargs):
        raise RuntimeError(self.reason)


def get_llm():
    """Build an LLM backend from current session-state settings."""
    provider = "ollama" if FORCE_OLLAMA_FOR_ALL_USERS else st.session_state.llm_provider

    if provider == "ollama":
        return LLMBackend(
            provider="ollama",
            model=st.session_state.llm_model,
            base_url=st.session_state.llm_base_url,
        )

    if provider == "hf_local":
        model_path = str(st.session_state.get("hf_local_model_path", "") or "").strip()
        if not model_path:
            reason = "HF Local model path is empty. Set it in Settings -> LLM Provider."
            st.session_state.hf_local_error = reason
            return _UnavailableLLM(reason)

        max_new_tokens_raw = st.session_state.get("hf_local_max_new_tokens", "2048")
        try:
            max_new_tokens = int(max_new_tokens_raw)
        except Exception:
            max_new_tokens = 2048

        try:
            from local_docgen_hf.integrate_existing_model import LocalHFBackend

            st.session_state.hf_local_error = ""
            return LocalHFBackend(model_path=model_path, default_max_tokens=max_new_tokens)
        except Exception as exc:
            reason = f"HF Local backend init failed: {exc}"
            st.session_state.hf_local_error = reason
            return _UnavailableLLM(reason)

    return LLMBackend(
        provider="openai",
        model=st.session_state.llm_model,
        api_key=st.session_state.openai_api_key,
    )


def get_sec_client() -> SECEdgarClient:
    return SECEdgarClient(user_agent=st.session_state.sec_user_agent)


# ======================================================================
# File processing
# ======================================================================

def _extract_legacy_doc_text(raw_bytes: bytes) -> str:
    """Best-effort .doc extraction using local Microsoft Word automation on Windows."""
    if not raw_bytes:
        return ""

    try:
        import win32com.client  # type: ignore
    except Exception:
        return ""

    temp_doc = None
    temp_txt = None
    word = None
    document = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tf:
            tf.write(raw_bytes)
            temp_doc = tf.name
        temp_txt = f"{temp_doc}.txt"

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(temp_doc, ReadOnly=True)
        wd_format_text = 2
        document.SaveAs(temp_txt, FileFormat=wd_format_text)

        with open(temp_txt, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception:
        return ""
    finally:
        try:
            if document is not None:
                document.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        for p in [temp_txt, temp_doc]:
            if p:
                try:
                    os.remove(p)
                except Exception:
                    pass


def _extract_google_doc_id(url: str) -> str:
    """Extract a Google Docs document id from common URL patterns."""
    parsed = urlparse((url or "").strip())
    if parsed.scheme not in {"http", "https"}:
        return ""

    host = (parsed.netloc or "").lower()
    if host not in {"docs.google.com", "www.docs.google.com"}:
        return ""

    parts = [p for p in (parsed.path or "").split("/") if p]
    # Expected path: /document/d/<doc_id>/...
    if len(parts) >= 3 and parts[0] == "document" and parts[1] == "d":
        return parts[2].strip()

    return ""


def extract_text_from_google_doc_url(url: str) -> str:
    """
    Fetch plain text from a shareable Google Doc URL.
    Requires the doc to be accessible to the current request context
    (for most local use-cases: "Anyone with the link" viewer access).
    """
    try:
        doc_id = _extract_google_doc_id(url)
        if not doc_id:
            return (
                "Error reading file: invalid Google Docs URL. "
                "Use a link like https://docs.google.com/document/d/<DOC_ID>/edit"
            )

        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"

        import requests
        response = requests.get(
            export_url,
            timeout=15,
            allow_redirects=True,
            headers={"User-Agent": "DocumentKnowledgeBase/1.0"},
        )

        if response.status_code in {401, 403}:
            return (
                "Error reading file: Google Doc access denied. "
                "Set sharing to 'Anyone with the link' (viewer) or use an accessible document."
            )

        if response.status_code >= 400:
            return f"Error reading file: Google Docs returned HTTP {response.status_code}."

        content_type = (response.headers.get("Content-Type") or "").lower()
        if "text/plain" not in content_type:
            return (
                "Error reading file: Google Docs did not return plain text. "
                "Ensure this is a Google Doc URL and sharing is configured correctly."
            )

        text_out = (response.text or "").replace("\r\n", "\n").strip("\ufeff\n ")
        if not text_out:
            return "Error reading file: imported Google Doc appears empty."
        return text_out
    except Exception as e:
        return f"Error reading file: {e}"


def extract_text_from_file(uploaded_file, filename: str | None = None) -> str:
    """Extract text from an uploaded .txt, .pdf, .doc, .docx, or .docm file."""
    try:
        name = (filename or getattr(uploaded_file, "name", "")).lower()
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)

        raw = uploaded_file.read() if hasattr(uploaded_file, "read") else uploaded_file
        if not isinstance(raw, (bytes, bytearray)):
            raw = str(raw).encode("utf-8", errors="ignore")

        if name.endswith(".txt"):
            return raw.decode("utf-8", errors="ignore")

        if name.endswith(".pdf"):
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        if name.endswith(".docx") or name.endswith(".docm"):
            import docx
            doc = docx.Document(io.BytesIO(raw))
            parts = []

            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)

            for table in doc.tables:
                for row in table.rows:
                    cells = [((c.text or "").strip()) for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))

            for section in doc.sections:
                for p in section.header.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)
                for p in section.footer.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)

            return "\n".join(parts)

        if name.endswith(".doc"):
            legacy_text = _extract_legacy_doc_text(bytes(raw))
            if legacy_text:
                return legacy_text
            return (
                "Error reading file: legacy .doc extraction requires Microsoft Word with pywin32 on this host. "
                "Please upload .docx for highest fidelity."
            )

        return "Unsupported file type."
    except Exception as e:
        return f"Error reading file: {e}"


def build_diff_highlight_html(original_text: str, revised_text: str) -> tuple[str, int, int]:
    """Return inline HTML diff and non-whitespace token counts (adds, deletes)."""

    def _tokenize(text: str) -> list[str]:
        return re.findall(r"\s+|[A-Za-z0-9_]+|[^\w\s]", text or "")

    old_tokens = _tokenize(original_text)
    new_tokens = _tokenize(revised_text)
    sm = difflib.SequenceMatcher(a=old_tokens, b=new_tokens)

    html_parts: list[str] = []
    add_count = 0
    del_count = 0

    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            html_parts.append(html_lib.escape("".join(new_tokens[j1:j2])))
        elif op == "insert":
            inserted = "".join(new_tokens[j1:j2])
            add_count += sum(1 for t in new_tokens[j1:j2] if t.strip())
            html_parts.append(
                f"<span style='color:#0b57d0;text-decoration:underline;text-decoration-thickness:2px;text-underline-offset:2px;padding:0 0.04rem;border-radius:2px;'>{html_lib.escape(inserted)}</span>"
            )
        elif op == "delete":
            deleted = "".join(old_tokens[i1:i2])
            del_count += sum(1 for t in old_tokens[i1:i2] if t.strip())
            html_parts.append(
                f"<span style='color:#b42318;text-decoration:line-through;text-decoration-thickness:2px;padding:0 0.04rem;border-radius:2px;'>{html_lib.escape(deleted)}</span>"
            )
        elif op == "replace":
            deleted = "".join(old_tokens[i1:i2])
            inserted = "".join(new_tokens[j1:j2])
            del_count += sum(1 for t in old_tokens[i1:i2] if t.strip())
            add_count += sum(1 for t in new_tokens[j1:j2] if t.strip())
            html_parts.append(
                f"<span style='color:#b42318;text-decoration:line-through;text-decoration-thickness:2px;padding:0 0.04rem;border-radius:2px;'>{html_lib.escape(deleted)}</span>"
            )
            html_parts.append(
                f"<span style='color:#0b57d0;text-decoration:underline;text-decoration-thickness:2px;text-underline-offset:2px;padding:0 0.04rem;border-radius:2px;'>{html_lib.escape(inserted)}</span>"
            )

    return "".join(html_parts), add_count, del_count


def build_diff_edit_markup(original_text: str, revised_text: str) -> tuple[str, int, int]:
    """Return editable redline markup using [[+added+]] and [[-deleted-]] markers."""

    def _tokenize(content: str) -> list[str]:
        return re.findall(r"\s+|[A-Za-z0-9_]+|[^\w\s]", content or "")

    old_tokens = _tokenize(original_text)
    new_tokens = _tokenize(revised_text)
    sm = difflib.SequenceMatcher(a=old_tokens, b=new_tokens)

    parts: list[str] = []
    add_count = 0
    del_count = 0

    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            parts.append("".join(new_tokens[j1:j2]))
        elif op == "insert":
            inserted = "".join(new_tokens[j1:j2])
            add_count += sum(1 for t in new_tokens[j1:j2] if t.strip())
            parts.append(f"[[+{inserted}+]]")
        elif op == "delete":
            deleted = "".join(old_tokens[i1:i2])
            del_count += sum(1 for t in old_tokens[i1:i2] if t.strip())
            parts.append(f"[[-{deleted}-]]")
        elif op == "replace":
            deleted = "".join(old_tokens[i1:i2])
            inserted = "".join(new_tokens[j1:j2])
            del_count += sum(1 for t in old_tokens[i1:i2] if t.strip())
            add_count += sum(1 for t in new_tokens[j1:j2] if t.strip())
            parts.append(f"[[-{deleted}-]][[+{inserted}+]]")

    return "".join(parts), add_count, del_count


def apply_edit_markup_to_text(markup_text: str) -> str:
    """Convert editable redline markup back to plain revised text."""
    if markup_text.count("[[+") != markup_text.count("+]]"):
        raise ValueError("Unbalanced addition markers. Use [[+added text+]].")
    if markup_text.count("[[-") != markup_text.count("-]]"):
        raise ValueError("Unbalanced deletion markers. Use [[-deleted text-]].")

    pattern = re.compile(r"\[\[\+(.*?)\+\]\]|\[\[\-(.*?)\-\]\]", flags=re.DOTALL)
    out = []
    cursor = 0
    for m in pattern.finditer(markup_text):
        out.append(markup_text[cursor:m.start()])
        added = m.group(1)
        deleted = m.group(2)
        if added is not None:
            out.append(added)
        elif deleted is not None:
            pass
        cursor = m.end()
    out.append(markup_text[cursor:])
    return "".join(out)


def _normalize_similarity_text(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").strip().lower())
    return re.sub(r"[^a-z0-9\s]", "", cleaned)


def _split_for_clause_comparison(text: str, max_units: int = 24, max_chars: int = 1200) -> list[dict]:
    text = (text or "").strip()
    if not text:
        return []

    units: list[dict] = []
    sections = [s.strip() for s in re.split(r"\n\s*\n", text) if s.strip()]
    for idx, sec in enumerate(sections, start=1):
        if len(sec) < 120:
            continue
        heading = sec.split("\n", 1)[0].strip()[:80]
        if len(heading) < 6:
            heading = f"Clause {idx}"
        if len(sec) > max_chars:
            sec = sec[:max_chars].rsplit(" ", 1)[0] + " ..."
        units.append({"title": heading, "text": sec})
        if len(units) >= max_units:
            break

    if not units:
        collapsed = re.sub(r"\s+", " ", text)
        chunk_size = max(800, min(max_chars, 1400))
        start = 0
        i = 1
        while start < len(collapsed) and len(units) < max_units:
            chunk = collapsed[start:start + chunk_size].strip()
            if chunk:
                units.append({"title": f"Clause {i}", "text": chunk})
            start += chunk_size
            i += 1

    return units


def _clip_text(text: str, limit: int = 700) -> str:
    text = (text or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit].rsplit(" ", 1)[0] + " ..."


def run_edgar_clause_compare(
    document_text: str,
    sec_client: SECEdgarClient,
    llm: LLMBackend,
    query: str,
    form_types: list[str],
    max_results: int = 4,
    similarity_floor: float = 0.34,
    max_suggestions: int = 5,
) -> dict:
    if not (query or "").strip():
        return {"error": "Enter an EDGAR query before running comparison."}

    hits = sec_client.search_filings(query=query.strip(), form_types=form_types or None, max_results=max_results)
    if not hits:
        return {"error": "No EDGAR filings matched that query."}

    user_units = _split_for_clause_comparison(document_text, max_units=20, max_chars=1200)
    if not user_units:
        return {"error": "Could not identify comparable clauses in the uploaded document."}

    peer_units: list[dict] = []
    used_hits: list[dict] = []
    retrieved_documents: list[dict] = []
    for hit in hits:
        filing_url = hit.get("url", "")
        if not filing_url:
            continue
        try:
            filing_text = sec_client.download_filing_text(filing_url, max_chars=60_000)
        except Exception:
            continue
        if len(filing_text) < 200:
            continue
        used_hits.append(hit)
        retrieved_documents.append({
            "entity_name": hit.get("entity_name", ""),
            "form_type": hit.get("form_type", ""),
            "file_date": hit.get("file_date", ""),
            "url": hit.get("url", ""),
            "text": _clip_text(filing_text, 14000),
        })
        for unit in _split_for_clause_comparison(filing_text, max_units=16, max_chars=1000):
            peer_units.append({"text": unit["text"], "title": unit["title"], "source": hit})

    if not peer_units:
        return {
            "error": "Matched filings were found, but no usable text could be extracted from them.",
            "documents": retrieved_documents,
        }

    candidates = []
    for u in user_units:
        user_norm = _normalize_similarity_text(u["text"])
        if len(user_norm) < 80:
            continue
        best = None
        for p in peer_units:
            peer_norm = _normalize_similarity_text(p["text"])
            if len(peer_norm) < 80:
                continue
            score = difflib.SequenceMatcher(None, user_norm, peer_norm).ratio()
            if best is None or score > best["score"]:
                best = {"score": score, "user_title": u["title"], "user_text": u["text"], "peer_title": p["title"], "peer_text": p["text"], "source": p["source"]}
        if best and best["score"] >= similarity_floor:
            candidates.append(best)

    if not candidates:
        return {
            "error": "No sufficiently similar clauses were found in EDGAR results. Try broader query or lower threshold.",
            "documents": retrieved_documents,
        }

    candidates.sort(key=lambda x: x["score"], reverse=True)
    suggestions = []
    can_use_llm = bool(llm and llm.is_available())

    for c in candidates[:max_suggestions]:
        source = c["source"]
        revision_prompt = (
            "Revise the clause below using the EDGAR peer benchmark while preserving this document's parties and structure.\n\n"
            f"Current clause:\n{_clip_text(c['user_text'], 900)}\n\n"
            f"Peer filing benchmark ({source.get('entity_name', 'Unknown')} {source.get('form_type', '')} {source.get('file_date', '')}):\n"
            f"{_clip_text(c['peer_text'], 900)}\n\n"
            "Integrate stronger language where appropriate and return the full updated document."
        )

        suggestion_text = ""
        if can_use_llm:
            try:
                suggestion_text = llm.chat([
                    {"role": "system", "content": "You are a legal editing assistant. Draft one concise clause-improvement recommendation grounded in a peer SEC filing."},
                    {"role": "user", "content": f"Current clause:\n{_clip_text(c['user_text'], 750)}\n\nPeer clause:\n{_clip_text(c['peer_text'], 750)}\n\nProvide: (1) what to change, (2) a short sample revision sentence. Keep under 120 words."},
                ], temperature=0.2, max_tokens=220).strip()
            except Exception:
                suggestion_text = ""
        if not suggestion_text:
            suggestion_text = "Align this clause with the cited peer filing by tightening defined triggers, adding objective standards, and preserving your party-specific terms."

        suggestions.append({
            "similarity": round(c["score"], 3),
            "current_title": c["user_title"],
            "peer_title": c["peer_title"],
            "current_excerpt": _clip_text(c["user_text"], 500),
            "peer_excerpt": _clip_text(c["peer_text"], 500),
            "suggestion_text": suggestion_text,
            "revision_prompt": revision_prompt,
            "source": {"entity_name": source.get("entity_name", ""), "form_type": source.get("form_type", ""), "file_date": source.get("file_date", ""), "url": source.get("url", "")},
        })

    return {
        "hits_considered": len(hits),
        "hits_used": len(used_hits),
        "peer_units": len(peer_units),
        "suggestions": suggestions,
        "documents": retrieved_documents,
    }



def search_edgar_comparable_documents(
    sec_client: SECEdgarClient,
    query: str,
    form_types: list[str],
    max_results: int = 10,
) -> dict:
    """Search comparable EDGAR filings and return top result metadata list."""
    if not (query or "").strip():
        return {"error": "Enter an EDGAR query to find comparable documents.", "results": []}

    hits = sec_client.search_filings(
        query=query.strip(),
        form_types=form_types or None,
        max_results=max_results,
    )
    if not hits:
        return {"error": "No comparable EDGAR documents found for that query.", "results": []}

    results = []
    for hit in hits:
        filing_url = hit.get("url", "")
        if not filing_url:
            continue
        results.append({
            "entity_name": hit.get("entity_name", ""),
            "form_type": hit.get("form_type", ""),
            "file_date": hit.get("file_date", ""),
            "url": filing_url,
        })

    if not results:
        return {"error": "Comparable filings were found, but no valid filing URLs were available.", "results": []}

    return {
        "results": results[:10],
        "hits_considered": len(hits),
        "hits_listed": min(10, len(results)),
    }


def load_edgar_document_preview(sec_client: SECEdgarClient, filing: dict) -> dict | None:
    """Load and normalize one EDGAR filing into preview text."""
    filing_url = filing.get("url", "")
    if not filing_url:
        return None
    try:
        filing_text = sec_client.download_filing_text(filing_url, max_chars=60_000)
    except Exception:
        return None
    if len(filing_text) < 200:
        return None
    return {
        "entity_name": filing.get("entity_name", ""),
        "form_type": filing.get("form_type", ""),
        "file_date": filing.get("file_date", ""),
        "url": filing_url,
        "text": _clip_text(filing_text, 14000),
    }


def infer_document_type_from_description(description: str) -> tuple[str, str]:
    """Infer best document type key from a natural language description."""
    desc = (description or "").strip()
    if not desc:
        return "custom_document", "No description provided"

    option_lines = [f"- {k}: {v.get('label', k)}" for k, v in DOCUMENT_TYPES.items()]
    option_block = "\n".join(option_lines)

    try:
        llm = get_llm()
        if llm.is_available():
            messages = [
                {
                    "role": "system",
                    "content": (
                        "You classify legal drafting requests. Return only one key from the list "
                        "or custom_document if uncertain."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Available keys:\n{option_block}\n\n"
                        f"Request description:\n{desc}\n\n"
                        "Return only the best key."
                    ),
                },
            ]
            raw = llm.chat(messages, temperature=0.0, max_tokens=32).strip().lower()
            key = raw.split()[0].strip("`.,:;\"' ") if raw else ""
            if key in DOCUMENT_TYPES:
                return key, "Detected by model"
            if key == "custom_document":
                return key, "Model requested custom flow"
    except Exception:
        pass

    d = desc.lower()
    fallback_rules = [
        ("contract", ["agreement", "contract", "vendor", "services", "msa", "sow"]),
        ("memo", ["memo", "memorandum", "internal analysis", "legal analysis"]),
        ("brief", ["brief", "motion", "argument", "court filing"]),
        ("corporate_filing", ["incorporation", "annual report", "amendment", "filing", "sec"]),
        ("nda", ["nda", "non-disclosure", "confidentiality"]),
        ("employment", ["employment", "employee", "offer letter", "termination"]),
        ("lease", ["lease", "rent", "landlord", "tenant"]),
    ]

    for key, words in fallback_rules:
        if key in DOCUMENT_TYPES and any(w in d for w in words):
            return key, "Detected by keyword rules"

    return "custom_document", "Defaulted to custom flow"


def render_workflow_header(title: str, subtitle: str, *, progress: float | None = None, step_note: str | None = None):
    """Render a consistent workflow header with optional progress + step note."""
    safe_title = html_lib.escape(title or "")
    safe_subtitle = html_lib.escape(subtitle or "")
    st.markdown(
        f"""
        <div class="workspace-hero">
            <h2>{safe_title}</h2>
            <p>{safe_subtitle}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if step_note:
        st.caption(step_note)
    if progress is not None:
        st.progress(progress)


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks


def process_uploaded_files(uploaded_files: List, document_type: str):
    """Extract, chunk, and store uploaded files with document_type metadata."""
    collection = get_collection()
    progress = st.progress(0)
    status = st.empty()
    total_chunks = 0
    new_files = 0
    skipped_files = 0
    failed_files = 0

    for i, f in enumerate(uploaded_files):
        filename = f.name
        user_id = get_data_scope_user_id()

        status.text(f"Processing {filename}...")
        try:
            text = extract_text_from_file(f)
            if text.startswith("Error reading file:"):
                st.error(text)
                failed_files += 1
                progress.progress((i + 1) / len(uploaded_files))
                continue

            content_hash = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()[:16]
            file_key = f"{user_id}_{filename}_{content_hash}"

            if file_key in st.session_state.processed_files:
                skipped_files += 1
                progress.progress((i + 1) / len(uploaded_files))
                continue

            chunks = chunk_text(text)
            if chunks:
                ids = [f"u{user_id}_{content_hash}_chunk_{j}" for j in range(len(chunks))]
                metadatas = [
                    {
                        "source": filename,
                        "chunk_index": j,
                        "document_type": document_type,
                        "user_id": user_id,
                        "content_hash": content_hash,
                    }
                    for j in range(len(chunks))
                ]

                try:
                    collection.delete(where={"$and": [{"source": filename}, {"user_id": user_id}]})
                except Exception:
                    pass

                collection.upsert(documents=chunks, ids=ids, metadatas=metadatas)
                total_chunks += len(chunks)
                new_files += 1
                st.session_state.processed_files.add(file_key)
            progress.progress((i + 1) / len(uploaded_files))
        except Exception as e:
            st.error(f"Error processing {filename}: {e}")
            failed_files += 1

    if new_files:
        st.toast(f"Processed {new_files} file(s) into {total_chunks} chunks.")
        st.success(f"Processed {new_files} file(s) into {total_chunks} chunks.")
    elif uploaded_files:
        st.info("All files have already been processed.")

    progress.empty()
    status.empty()
    return {
        "new_files": new_files,
        "total_chunks": total_chunks,
        "skipped_files": skipped_files,
        "failed_files": failed_files,
    }


# ======================================================================
# RAG query (Chat tab)
# ======================================================================

def rag_query(question: str) -> tuple[str, list[str]]:
    """Retrieve relevant chunks and generate an answer via the configured LLM."""
    collection = get_collection()
    if collection.count() == 0:
        return "No documents indexed yet. Upload and process documents first.", []

    results = collection.query(
        query_texts=[question],
        n_results=min(5, collection.count()),
    )
    documents = results["documents"][0]
    metadatas = results["metadatas"][0]

    context_parts = []
    sources = set()
    for doc, meta in zip(documents, metadatas):
        source = meta["source"]
        sources.add(source)
        context_parts.append(f"[Source: {source}]\n{doc}")
    context = "\n\n---\n\n".join(context_parts)

    messages = [
        {
            "role": "system",
            "content": (
                "You are a helpful legal research assistant. Answer the user's "
                "question based on the provided document context. Cite which "
                "source document(s) the information comes from. If the context "
                "doesn't contain enough information, say so clearly."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Context from uploaded documents:\n{context}\n\n"
                f"Question: {question}\n\n"
                "Provide a helpful answer citing source document(s):"
            ),
        },
    ]

    try:
        llm = get_llm()
        answer = llm.chat(messages, temperature=0.1, max_tokens=1024)
        return answer, sorted(sources)
    except Exception as e:
        return f"Error generating answer: {e}", []


# ======================================================================
# Database stats helper
# ======================================================================

def get_db_stats() -> dict:
    """Return chunk counts broken down by document_type."""
    collection = get_collection()
    total = collection.count()
    stats = {"total": total}
    if total == 0:
        return stats

    for dtype in list(DOCUMENT_TYPES.keys()) + ["reference"]:
        try:
            qr = collection.query(
                query_texts=["document"],
                n_results=min(total, 1000),
                where={"document_type": dtype},
                include=[],
            )
            stats[dtype] = len(qr["ids"][0]) if qr["ids"] else 0
        except Exception:
            stats[dtype] = 0
    return stats


# ======================================================================
# Model improvement job control (UI-managed background processes)
# ======================================================================

PROJECT_ROOT = Path(__file__).resolve().parent
CONTROL_DIR = PROJECT_ROOT / "artifacts" / "ui_jobs"
CONTROL_DIR.mkdir(parents=True, exist_ok=True)


def _job_files(job_key: str) -> tuple[Path, Path, Path]:
    safe_key = re.sub(r"[^a-zA-Z0-9_-]+", "_", job_key).lower()
    return (
        CONTROL_DIR / f"{safe_key}.pid",
        CONTROL_DIR / f"{safe_key}.json",
        CONTROL_DIR / f"{safe_key}.log",
    )


def _is_pid_running(pid: int) -> bool:
    if pid <= 0:
        return False
    result = subprocess.run(
        ["tasklist", "/FI", f"PID eq {pid}"],
        capture_output=True,
        text=True,
    )
    return str(pid) in (result.stdout or "")


def get_job_status(job_key: str) -> dict:
    pid_file, meta_file, log_file = _job_files(job_key)
    status = {"running": False, "pid": None, "meta": {}, "log_file": str(log_file)}

    if meta_file.exists():
        try:
            status["meta"] = json.loads(meta_file.read_text(encoding="utf-8"))
        except Exception:
            status["meta"] = {}

    if pid_file.exists():
        try:
            pid = int(pid_file.read_text(encoding="utf-8").strip())
            status["pid"] = pid
            status["running"] = _is_pid_running(pid)
        except Exception:
            status["running"] = False

    if not status["running"] and pid_file.exists():
        try:
            pid_file.unlink()
        except Exception:
            pass

    return status


def start_job(job_key: str, command: list[str], env_overrides: dict[str, str] | None = None) -> tuple[bool, str]:
    pid_file, meta_file, log_file = _job_files(job_key)
    status = get_job_status(job_key)
    if status["running"]:
        return False, f"Job already running (PID {status['pid']})."

    log_file.parent.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    # Force immediate log flushing from Python child processes.
    env["PYTHONUNBUFFERED"] = "1"
    if env_overrides:
        env.update({k: v for k, v in env_overrides.items() if v is not None})

    creation_flags = 0
    if hasattr(subprocess, "DETACHED_PROCESS"):
        creation_flags |= subprocess.DETACHED_PROCESS
    if hasattr(subprocess, "CREATE_NEW_PROCESS_GROUP"):
        creation_flags |= subprocess.CREATE_NEW_PROCESS_GROUP

    with open(log_file, "a", encoding="utf-8") as log_handle:
        log_handle.write(f"\n\n=== START {job_key} ===\n")
        log_handle.write("COMMAND: " + " ".join(command) + "\n")
        proc = subprocess.Popen(
            command,
            cwd=str(PROJECT_ROOT),
            stdout=log_handle,
            stderr=subprocess.STDOUT,
            env=env,
            creationflags=creation_flags,
        )

    pid_file.write_text(str(proc.pid), encoding="utf-8")
    meta = {
        "started_at": __import__("datetime").datetime.now().isoformat(),
        "command": command,
        "env_keys": sorted(list((env_overrides or {}).keys())),
        "log_file": str(log_file),
    }
    meta_file.write_text(json.dumps(meta, indent=2), encoding="utf-8")
    return True, f"Started {job_key} (PID {proc.pid})."


def stop_job(job_key: str) -> tuple[bool, str]:
    pid_file, _meta_file, _log_file = _job_files(job_key)
    status = get_job_status(job_key)
    pid = status.get("pid")
    if not pid:
        return False, "No PID found for this job."
    if not status["running"]:
        return False, "Job is not running."

    result = subprocess.run(["taskkill", "/PID", str(pid), "/T", "/F"], capture_output=True, text=True)
    if result.returncode == 0:
        try:
            pid_file.unlink()
        except Exception:
            pass
        return True, f"Stopped {job_key} (PID {pid})."

    return False, (result.stderr or result.stdout or "Failed to stop job").strip()


def read_job_log(job_key: str, max_lines: int = 120) -> str:
    _pid_file, _meta_file, log_file = _job_files(job_key)
    if not log_file.exists():
        return ""
    try:
        lines = log_file.read_text(encoding="utf-8", errors="ignore").splitlines()
        return "\n".join(lines[-max_lines:])
    except Exception:
        return ""


def list_job_keys(prefix: str) -> list[str]:
    normalized = re.sub(r"[^a-zA-Z0-9_-]+", "_", prefix).lower()
    discovered: set[str] = set()
    for ext in ("json", "pid", "log"):
        for path in CONTROL_DIR.glob(f"{normalized}*.{ext}"):
            discovered.add(path.stem)
    return sorted(discovered, reverse=True)


def _command_option_value(command: list[str], option_name: str) -> str | None:
    for idx, token in enumerate(command):
        if token == option_name and idx + 1 < len(command):
            return command[idx + 1]
    return None


def _format_duration(seconds: float | int | None) -> str:
    if seconds is None:
        return "-"
    total = int(max(0, float(seconds)))
    hours, rem = divmod(total, 3600)
    minutes, secs = divmod(rem, 60)
    if hours:
        return f"{hours}h {minutes}m {secs}s"
    if minutes:
        return f"{minutes}m {secs}s"
    return f"{secs}s"


def _canonical_edgar_queries() -> list[str]:
    return [
        "material agreement",
        "credit agreement",
        "merger agreement",
        "employment agreement",
        "risk factors",
        "securities offering",
        "corporate governance",
        "litigation",
        "proxy statement",
        "related party transactions",
    ]


def derive_edgar_queries_from_prompt(base_prompt: str, max_queries: int = 5) -> str:
    prompt = (base_prompt or "").lower()
    query_candidates: list[str] = []

    keyword_map = [
        (("merger", "acquisition", "m&a"), "merger agreement"),
        (("credit", "loan", "debt", "facility"), "credit agreement"),
        (("employment", "executive", "compensation"), "employment agreement"),
        (("risk", "uncertainty"), "risk factors"),
        (("governance", "board", "director"), "corporate governance"),
        (("litigation", "dispute", "claim"), "litigation"),
        (("offering", "security", "capital raise"), "securities offering"),
        (("proxy", "shareholder vote"), "proxy statement"),
        (("related party", "affiliate"), "related party transactions"),
        (("agreement", "contract", "obligation"), "material agreement"),
    ]

    for keywords, query in keyword_map:
        if any(k in prompt for k in keywords):
            query_candidates.append(query)

    if not query_candidates:
        query_candidates = _canonical_edgar_queries()[:max_queries]

    deduped: list[str] = []
    for query in query_candidates:
        if query not in deduped:
            deduped.append(query)

    return ",".join(deduped[:max_queries])


def combine_edgar_query_inputs(*query_sets: str) -> str:
    seen: set[str] = set()
    merged: list[str] = []
    for raw in query_sets:
        for part in (raw or "").split(","):
            query = part.strip()
            if query and query.lower() not in seen:
                seen.add(query.lower())
                merged.append(query)
    return ",".join(merged)


def _extract_json_array(text: str) -> list:
    raw = (text or "").strip()
    if not raw:
        return []
    try:
        value = json.loads(raw)
        return value if isinstance(value, list) else []
    except Exception:
        pass
    start = raw.find("[")
    end = raw.rfind("]")
    if start >= 0 and end > start:
        try:
            value = json.loads(raw[start : end + 1])
            return value if isinstance(value, list) else []
        except Exception:
            return []
    return []


def generate_edgar_queries_with_local_model(base_prompt: str, max_queries: int = 8) -> tuple[str, str]:
    prompt = (base_prompt or "").strip()
    if not prompt:
        return "", "Enter a learning objective first."

    try:
        local_llm = LLMBackend(
            provider="ollama",
            model=st.session_state.llm_model,
            base_url=st.session_state.llm_base_url,
        )
        messages = [
            {"role": "system", "content": "Return strict JSON array only."},
            {
                "role": "user",
                "content": (
                    "Given this legal learning objective, generate targeted SEC EDGAR filing search queries. "
                    f"Return JSON array of up to {max_queries} short query strings only (no objects). "
                    "Focus on filing topics likely to improve legal drafting and factual extraction.\n\n"
                    f"Learning objective:\n{prompt}"
                ),
            },
        ]
        raw = local_llm.chat(messages, temperature=0.1, top_p=0.9, top_k=40, max_tokens=500)
        items = _extract_json_array(raw)
        cleaned = [
            str(x).strip()
            for x in items
            if isinstance(x, str) and str(x).strip()
        ]
        if cleaned:
            return combine_edgar_query_inputs(",".join(cleaned[:max_queries])), "Generated by local model."
        return "", "Local model returned no usable EDGAR queries. Adjust prompt and try again."
    except Exception as exc:
        return "", f"Local model query generation failed: {exc}"


def read_strategy_progress(strategy_status: dict) -> dict | None:
    progress_path = PROJECT_ROOT / "artifacts" / "agent_improvement" / "progress.json"
    command = strategy_status.get("meta", {}).get("command", [])
    if isinstance(command, list):
        output_dir = _command_option_value(command, "--output-dir")
        if output_dir:
            progress_path = PROJECT_ROOT / output_dir / "progress.json"

    if not progress_path.exists():
        return None

    try:
        return json.loads(progress_path.read_text(encoding="utf-8"))
    except Exception:
        return None


def _new_strategy_job_key() -> str:
    return f"strategy_agents_{datetime.utcnow().strftime('%Y%m%d_%H%M%S_%f')}"


def render_model_improvement_page():
    """Admin UI to run/monitor model-improvement jobs without Task Scheduler."""
    current_user = st.session_state.current_user
    if not current_user.is_admin():
        st.error("Access denied. Model improvement controls are admin-only.")
        return

    if st.button("Back to Workspace", key="back_from_model_improvement"):
        st.session_state.show_model_improvement = False
        st.rerun()

    render_workflow_header(
        "Model Improvement",
        "Start/stop background jobs for strategy optimization and true weight training.",
        step_note="You can control both jobs from this page and adjust learning preferences.",
    )
    st.markdown(
        """
        <style>
        div[data-testid="stButton"] button[kind="primary"] {
            background-color: #0f766e !important;
            color: #ffffff !important;
            border: 1px solid #0f766e !important;
        }
        div[data-testid="stButton"] button[kind="primary"]:hover {
            background-color: #0d9488 !important;
            color: #ffffff !important;
            border: 1px solid #0d9488 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    strategy_tab, weight_tab = st.tabs(["Strategy Agents", "True Weight Training"])

    with strategy_tab:
        st.markdown("#### Strategy optimization (no weight changes)")
        mode = st.selectbox("Data mode", ["hybrid", "edgar", "uploads"], index=0, key="mi_strategy_mode")
        iterations = st.number_input("Iterations", min_value=1, max_value=50, value=6, key="mi_strategy_iterations")
        candidates = st.number_input("Candidates per iteration", min_value=2, max_value=20, value=8, key="mi_strategy_candidates")
        default_strategy_prompt = (
            "You are a legal analysis assistant. Use only provided context, cite concrete facts, "
            "and do not invent missing information."
        )
        if "mi_strategy_prompt" not in st.session_state:
            st.session_state.mi_strategy_prompt = default_strategy_prompt
        if "mi_strategy_edgar_additional_queries" not in st.session_state:
            st.session_state.mi_strategy_edgar_additional_queries = ""
        if "mi_strategy_generated_queries" not in st.session_state:
            st.session_state.mi_strategy_generated_queries = st.session_state.get("mi_strategy_edgar_base_queries", "")

        base_prompt = st.text_area(
            "Learning objective / base system prompt",
            height=120,
            key="mi_strategy_prompt",
        )

        if st.button(
            "Run Prompt to Generate EDGAR Queries (Local Model)",
            type="primary",
            use_container_width=True,
            key="mi_generate_edgar_queries",
        ):
            generated_queries, message = generate_edgar_queries_with_local_model(base_prompt, max_queries=8)
            st.session_state.mi_strategy_generated_queries = generated_queries
            if generated_queries:
                st.success(message)
            else:
                st.warning(message)

        generated_edgar_queries = st.text_input(
            "Generated EDGAR queries (editable, comma-separated)",
            key="mi_strategy_generated_queries",
        )
        additional_edgar_queries = st.text_input(
            "Additional EDGAR queries (optional, comma-separated)",
            key="mi_strategy_edgar_additional_queries",
        )
        effective_edgar_queries = combine_edgar_query_inputs(generated_edgar_queries, additional_edgar_queries)
        st.caption(f"Effective EDGAR queries: {effective_edgar_queries or '(none)'}")

        strategy_job_keys = list_job_keys("strategy_agents")
        strategy_jobs: list[dict] = []
        for key in strategy_job_keys:
            status = get_job_status(key)
            strategy_jobs.append({"job_key": key, **status})

        running_count = sum(1 for job in strategy_jobs if job.get("running"))
        st.caption(f"Strategy jobs: {len(strategy_jobs)} total, {running_count} running")

        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("Start Strategy Job", type="primary", use_container_width=True, key="mi_start_strategy"):
                if not effective_edgar_queries.strip():
                    st.warning("Generate EDGAR queries first, or enter queries manually before starting.")
                    st.stop()
                job_key = _new_strategy_job_key()
                output_dir = f"artifacts/agent_improvement/{job_key}"
                cmd = [
                    sys.executable,
                    "-u",
                    "-m",
                    "agents.multi_agent_improver",
                    "--mode",
                    mode,
                    "--iterations",
                    str(int(iterations)),
                    "--candidates-per-iteration",
                    str(int(candidates)),
                    "--edgar-queries",
                    effective_edgar_queries,
                    "--base-system-prompt",
                    base_prompt,
                    "--output-dir",
                    output_dir,
                ]
                ok, msg = start_job(job_key, cmd, env_overrides={"SEC_EDGAR_USER_AGENT": st.session_state.get("sec_user_agent", "")})
                if ok:
                    st.success(msg)
                else:
                    st.warning(msg)
        with c2:
            if st.button("Stop All Running Strategy Jobs", use_container_width=True, key="mi_stop_strategy_all"):
                running_jobs = [job for job in strategy_jobs if job.get("running")]
                stopped = 0
                failures = 0
                for job in running_jobs:
                    ok, _msg = stop_job(job["job_key"])
                    if ok:
                        stopped += 1
                    else:
                        failures += 1
                if stopped:
                    st.success(f"Stopped {stopped} strategy job(s).")
                if failures:
                    st.warning(f"Failed to stop {failures} strategy job(s).")
                if not running_jobs:
                    st.info("No running strategy jobs found.")
        with c3:
            if st.button("Refresh Status", use_container_width=True, key="mi_refresh_strategy"):
                st.rerun()

        if strategy_jobs:
            st.markdown("#### Strategy job monitor")
            for idx, job in enumerate(strategy_jobs):
                job_key = job["job_key"]
                running = bool(job.get("running"))
                status_label = "Running" if running else "Stopped"
                started_at = str(job.get("meta", {}).get("started_at") or "unknown")
                with st.expander(f"{job_key} | {status_label} | started {started_at}", expanded=(idx == 0)):
                    st.caption(f"PID: {job.get('pid') or '-'}")
                    progress = read_strategy_progress(job)
                    p1, p2, p3 = st.columns(3)
                    if progress:
                        stage = str(progress.get("stage") or "unknown").replace("_", " ").title()
                        percent = float(progress.get("percent") or 0.0)
                        eta_seconds = progress.get("eta_seconds")
                        with p1:
                            st.metric("Progress", f"{percent:.1f}%")
                        with p2:
                            st.metric("Stage", stage)
                        with p3:
                            st.metric("ETA", _format_duration(eta_seconds))
                        st.progress(max(0.0, min(1.0, percent / 100.0)))
                        if progress.get("message"):
                            st.caption(str(progress["message"]))
                        if progress.get("iteration") and progress.get("total_iterations"):
                            st.caption(f"Iteration: {progress['iteration']}/{progress['total_iterations']}")
                        if progress.get("best_score") is not None:
                            st.caption(f"Best score: {float(progress['best_score']):.4f}")
                    else:
                        with p1:
                            st.metric("Progress", "0.0%")
                        with p2:
                            st.metric("Stage", "Starting")
                        with p3:
                            st.metric("ETA", "-")
                        st.caption("No progress file yet.")

                    if running and st.button("Stop This Job", use_container_width=True, key=f"mi_stop_strategy_{job_key}"):
                        ok, msg = stop_job(job_key)
                        if ok:
                            st.success(msg)
                            st.rerun()
                        st.warning(msg)

                    strategy_log = read_job_log(job_key, max_lines=80)
                    st.text_area(
                        "Execution log (tail)",
                        value=strategy_log,
                        height=180,
                        key=f"mi_strategy_log_area_{job_key}",
                    )
        else:
            st.info("No strategy jobs yet. Start a job to see real-time progress here.")

    with weight_tab:
        st.markdown("#### True weight improvement (LoRA fine-tuning + Ollama export)")
        include_uploads = st.checkbox("Include uploads corpus", value=True, key="mi_weight_include_uploads")
        include_edgar = st.checkbox("Include EDGAR corpus", value=True, key="mi_weight_include_edgar")
        use_fallback = st.checkbox("Use fallback training settings (lower memory)", value=True, key="mi_weight_fallback")
        edgar_queries_w = st.text_input(
            "EDGAR queries (comma-separated)",
            value="material agreement,credit agreement,merger agreement,employment agreement,risk factors",
            key="mi_weight_edgar_queries",
        )
        doc_type_guidance = st.text_area(
            "Document-type guidance prompt",
            value="Prioritize corporate legal drafting templates with clear party, term, obligation, remedy, and governing-law fields.",
            height=100,
            key="mi_doc_type_guidance",
        )

        weight_status = get_job_status("weight_training")
        st.caption(f"Status: {'Running' if weight_status['running'] else 'Stopped'}")
        if weight_status["pid"]:
            st.caption(f"PID: {weight_status['pid']}")

        w1, w2, w3 = st.columns(3)
        with w1:
            if st.button("Start Weight Training", type="primary", use_container_width=True, key="mi_start_weight"):
                cmd = [
                    sys.executable,
                    "-u",
                    "finetune/continuous_weight_improvement.py",
                    "--edgar-queries",
                    edgar_queries_w,
                    "--doc-type-guidance",
                    doc_type_guidance,
                ]
                if include_uploads:
                    cmd.append("--include-uploads")
                if include_edgar:
                    cmd.append("--include-edgar")
                if use_fallback:
                    cmd.append("--fallback")
                ok, msg = start_job("weight_training", cmd, env_overrides={"SEC_EDGAR_USER_AGENT": st.session_state.get("sec_user_agent", "")})
                if ok:
                    st.success(msg)
                else:
                    st.warning(msg)
        with w2:
            if st.button("Stop Weight Training", use_container_width=True, key="mi_stop_weight"):
                ok, msg = stop_job("weight_training")
                if ok:
                    st.success(msg)
                else:
                    st.warning(msg)
        with w3:
            if st.button("Refresh Status ", use_container_width=True, key="mi_refresh_weight"):
                st.rerun()

        weight_log = read_job_log("weight_training", max_lines=120)
        st.text_area("Weight training log (tail)", value=weight_log, height=260, key="mi_weight_log_area")


# ======================================================================
# SIDEBAR
# ======================================================================

def render_sidebar():
    """Render generation-focused sidebar. Operational controls are admin-only."""
    current_user = st.session_state.current_user

    render_auth_sidebar(current_user)

    if current_user.is_admin():
        st.markdown("""
        <div class="card-header">
            Admin
        </div>
        """, unsafe_allow_html=True)

        if st.button("Settings", use_container_width=True, key="sidebar_settings"):
            st.session_state.show_settings = True
            st.session_state.show_knowledge_base = False
            st.session_state.show_model_improvement = False
            st.rerun()

        if st.button("Knowledge Base", use_container_width=True, key="sidebar_knowledge"):
            st.session_state.show_knowledge_base = True
            st.session_state.show_settings = False
            st.session_state.show_model_improvement = False
            st.rerun()

        if st.button("Model Improvement", use_container_width=True, key="sidebar_model_improvement"):
            st.session_state.show_model_improvement = True
            st.session_state.show_knowledge_base = False
            st.session_state.show_settings = False
            st.rerun()

        st.divider()

    llm = get_llm()
    if llm.is_available():
        provider_label = st.session_state.llm_provider.title()
        st.markdown(f"""
        <div class="status-badge status-success">
            <span class="connection-indicator connection-success"></span>
            {provider_label} ({st.session_state.llm_model})
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="status-badge status-error">
            <span class="connection-indicator connection-error"></span>
            LLM: Not connected
        </div>
        """, unsafe_allow_html=True)



# ======================================================================
# ONBOARDING WIZARD
# ======================================================================

def render_onboarding_wizard():
    """Render the first-run onboarding wizard for API key setup."""
    st.markdown("""
    <div style="max-width: 700px; margin: 2rem auto;">
        <div class="card fade-in">
            <div style="text-align: center; margin-bottom: 2rem;">
                <div style="font-size: 4rem; margin-bottom: 1rem;">ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â½ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°</div>
                <h1 style="margin-bottom: 0.5rem;">Welcome to Your Document Generator!</h1>
                <p style="color: var(--text-secondary); font-size: 1.1rem;">
                    Let's get you set up in just a minute
                </p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 3, 1])

    with col2:
        st.markdown("""
        <div class="card fade-in">
            <h3 style="margin-top: 0;">ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¹ What This App Does</h3>
            <ul style="line-height: 2; color: var(--text-secondary);">
                <li><strong>Upload Documents:</strong> Add your legal documents to build a knowledge base</li>
                <li><strong>Ask Questions:</strong> Chat with your documents using AI-powered search</li>
                <li><strong>Generate Documents:</strong> Create contracts, memos, briefs, and filings</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("""
        <div class="card fade-in">
            <h3 style="margin-top: 0;">ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¹Ãƒâ€¦Ã¢â‚¬Å“ Setup Your AI Provider</h3>
            <p style="color: var(--text-secondary);">
                This app needs an AI language model to function. Choose your preferred option:
            </p>
        </div>
        """, unsafe_allow_html=True)

        provider_choice = st.radio(
            "Select your AI provider:",
            options=["ollama", "openai"],
            format_func=lambda x: "Ollama (Recommended ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â free, runs locally)" if x == "ollama" else "OpenAI (Cloud, requires API key)",
            index=0,
            key="onboarding_provider"
        )

        if provider_choice == "openai":
            st.markdown("""
            <div class="info-card">
                <strong>OpenAI Setup:</strong><br>
                1. Visit <a href="https://platform.openai.com/api-keys" target="_blank">OpenAI Platform</a><br>
                2. Sign up or log in<br>
                3. Create a new API key<br>
                4. Copy and paste it below
            </div>
            """, unsafe_allow_html=True)

            api_key = st.text_input(
                "Enter your OpenAI API Key:",
                type="password",
                placeholder="sk-proj-...",
                help="Your API key should start with 'sk-'",
                key="onboarding_api_key"
            )

            if api_key:
                if api_key.startswith("sk-") and len(api_key) > 20:
                    st.markdown("""
                    <div class="success-card">
                        ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ API key format looks valid!
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div class="warning-card">
                        ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â API key format may be invalid (should start with 'sk-')
                    </div>
                    """, unsafe_allow_html=True)

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¾Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¾ Save & Continue", type="primary", use_container_width=True, disabled=not api_key):
                    st.session_state.openai_api_key = api_key
                    st.session_state.llm_provider = "openai"
                    st.session_state.llm_model = "gpt-4o-mini"
                    st.session_state.onboarding_complete = True
                    st.toast("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ OpenAI configured successfully!", icon="ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦")
                    st.rerun()

            with col_btn2:
                if st.button("Skip for Now", use_container_width=True):
                    st.session_state.onboarding_complete = True
                    st.rerun()

        else:  # Ollama
            st.markdown("""
            <div class="info-card">
                <strong>Ollama Setup:</strong><br>
                1. Download from <a href="https://ollama.com" target="_blank">ollama.com</a><br>
                2. Install and run Ollama on your computer<br>
                3. Pull a model: <code>ollama pull llama3.1:8b</code><br>
                4. Ollama should be running at http://localhost:11434
            </div>
            """, unsafe_allow_html=True)

            # Check if Ollama is available
            try:
                import requests
                r = requests.get("http://localhost:11434/api/tags", timeout=2)
                if r.status_code == 200:
                    st.markdown("""
                    <div class="success-card">
                        ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ Ollama detected and running!
                    </div>
                    """, unsafe_allow_html=True)

                    col_btn1, col_btn2 = st.columns(2)
                    with col_btn1:
                        if st.button("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¾Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¾ Use Ollama", type="primary", use_container_width=True):
                            st.session_state.llm_provider = "ollama"
                            st.session_state.llm_model = "llama3.1:8b"
                            st.session_state.onboarding_complete = True
                            st.toast("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ Ollama configured successfully!", icon="ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦")
                            st.rerun()
                    with col_btn2:
                        if st.button("Skip for Now", use_container_width=True):
                            st.session_state.onboarding_complete = True
                            st.rerun()
                else:
                    st.markdown("""
                    <div class="warning-card">
                        ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Ollama not detected. Please install and start Ollama first.
                    </div>
                    """, unsafe_allow_html=True)

                    if st.button("Skip for Now", use_container_width=True):
                        st.session_state.onboarding_complete = True
                        st.rerun()
            except Exception:
                st.markdown("""
                <div class="warning-card">
                    ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Ollama not detected. Please install and start Ollama first.
                </div>
                """, unsafe_allow_html=True)

                if st.button("Skip for Now", use_container_width=True):
                    st.session_state.onboarding_complete = True
                    st.rerun()


# ======================================================================
# TAB 1 ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Chat Q&A
# ======================================================================

def render_chat_tab():
    """Render the chat Q&A interface."""
    st.markdown("""
    <div class="card-header">
        ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¾Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ Chat with Your Documents
    </div>
    """, unsafe_allow_html=True)

    collection = get_collection()
    if collection.count() == 0:
        # Better empty state with onboarding
        st.markdown("""
        <div style="max-width: 600px; margin: 3rem auto; text-align: center;">
            <div style="font-size: 5rem; margin-bottom: 1.5rem; opacity: 0.6;">ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡</div>
            <h3 style="color: var(--text-primary); margin-bottom: 1rem;">
                No Documents Yet
            </h3>
            <p style="color: var(--text-secondary); font-size: 1.1rem; margin-bottom: 2rem;">
                Upload legal documents in the sidebar to start asking questions about them.
            </p>
        </div>
        """, unsafe_allow_html=True)

        # Quick guide
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("""
            <div class="card">
                <div style="font-size: 2.5rem; text-align: center; margin-bottom: 1rem;">1ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬ÃƒÂ¢Ã¢â‚¬Å¾Ã‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â£</div>
                <h4 style="text-align: center; margin-bottom: 0.5rem;">Upload</h4>
                <p style="text-align: center; color: var(--text-secondary); font-size: 0.9rem;">
                    Add PDF, DOCX, or TXT files using the sidebar
                </p>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown("""
            <div class="card">
                <div style="font-size: 2.5rem; text-align: center; margin-bottom: 1rem;">2ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬ÃƒÂ¢Ã¢â‚¬Å¾Ã‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â£</div>
                <h4 style="text-align: center; margin-bottom: 0.5rem;">Process</h4>
                <p style="text-align: center; color: var(--text-secondary); font-size: 0.9rem;">
                    Click "Process Documents" to index them
                </p>
            </div>
            """, unsafe_allow_html=True)
        with col3:
            st.markdown("""
            <div class="card">
                <div style="font-size: 2.5rem; text-align: center; margin-bottom: 1rem;">3ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬ÃƒÂ¢Ã¢â‚¬Å¾Ã‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â£</div>
                <h4 style="text-align: center; margin-bottom: 0.5rem;">Ask</h4>
                <p style="text-align: center; color: var(--text-secondary); font-size: 0.9rem;">
                    Ask questions and get AI-powered answers
                </p>
            </div>
            """, unsafe_allow_html=True)
        return

    st.caption("Ask questions about your uploaded documents. The AI will search your knowledge base and provide answers with sources.")

    # Chat history
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # Chat input
    if prompt := st.chat_input("Ask a question about your documents..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Searching documents and generating answer..."):
                response, sources = rag_query(prompt)
            st.markdown(response)
            if sources:
                st.caption(f"ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡ Sources: {', '.join(sources)}")
            full = response
            if sources:
                full += f"\n\n*Sources: {', '.join(sources)}*"
            st.session_state.messages.append({"role": "assistant", "content": full})

        # Clear chat button
        if len(st.session_state.messages) > 0:
            if st.button("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¹Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Clear Chat History"):
                st.session_state.messages = []
                st.rerun()


# ======================================================================
# TAB 2 ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Generate Document - Workflow Functions
# ======================================================================

def render_mimic_workflow():
    """Workflow 1: Upload and mimic a reference document."""
    st.markdown("""
    <div style="margin-bottom: 1rem;">
        <h3 style="margin-bottom: 0.5rem;">Upload a Reference Document</h3>
        <p style="color: var(--text-secondary);">
            Upload a document you'd like to replicate. The AI will analyze its structure,
            extract key fields, and let you modify values to create a new document in the same style.
        </p>
    </div>
    """, unsafe_allow_html=True)

    # File upload
    uploaded_ref = st.file_uploader(
        "Upload reference document",
        type=["pdf", "docx", "txt"],
        key="mimic_upload",
        help="Upload a document to use as a template"
    )

    if uploaded_ref:
        # Extract text
        with st.spinner("Reading document..."):
            ref_text = extract_text_from_file(uploaded_ref)

        if len(ref_text) < 100:
            st.error("Document appears to be too short or unreadable. Please upload a valid document.")
            return

        st.success(f"ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ Loaded {len(ref_text)} characters from {uploaded_ref.name}")

        # Analyze the document
        if st.button("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Analyze Document Structure", type="primary", use_container_width=True):
            with st.spinner("Analyzing document structure and extracting fields..."):
                llm = get_llm()
                collection = get_collection()
                generator = DocumentGenerator(llm, collection, knowledge_db=get_knowledge_db())
                analysis = generator.analyze_document(ref_text)

                # Store in session state
                st.session_state.mimic_analysis = analysis
                st.session_state.mimic_ref_text = ref_text
                st.toast("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ Analysis complete!", icon="ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦")
                st.rerun()

        # Show analysis results and editable fields
        if "mimic_analysis" in st.session_state:
            analysis = st.session_state.mimic_analysis

            st.divider()
            st.markdown(f"""
            <div class="success-card">
                <strong>ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦Ãƒâ€šÃ‚Â¾ Document Type:</strong> {analysis.get('document_subtype', 'Unknown')}<br>
                <strong>ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â½ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¨ Tone:</strong> {analysis.get('tone', 'formal')}<br>
                <strong>ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Style:</strong> {analysis.get('style_notes', 'Standard legal document')}
            </div>
            """, unsafe_allow_html=True)

            st.markdown("### ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Edit Field Values")
            st.caption("Modify the extracted values below. The AI will generate a new document with your changes.")

            # Editable form for extracted fields
            with st.form("mimic_edit_form"):
                user_edits = {}
                key_fields = analysis.get("key_fields", {})

                if not key_fields:
                    st.info("No specific fields were extracted. You can describe changes in the text area below.")
                    user_edits["general_changes"] = st.text_area(
                        "Describe the changes you want to make",
                        placeholder="E.g., Change party names from X to Y, update date to...",
                        height=150
                    )
                else:
                    for field_name, field_value in key_fields.items():
                        user_edits[field_name] = st.text_input(
                            field_name.replace("_", " ").title(),
                            value=str(field_value),
                            key=f"mimic_{field_name}"
                        )

                st.divider()
                submitted = st.form_submit_button(
                    "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¨ Generate Document",
                    type="primary",
                    use_container_width=True
                )

                if submitted:
                    with st.spinner("Generating document in the style of your reference..."):
                        llm = get_llm()
                        collection = get_collection()
                        generator = DocumentGenerator(llm, collection, knowledge_db=get_knowledge_db())

                        try:
                            text = generator.generate_from_template(
                                st.session_state.mimic_ref_text,
                                analysis,
                                user_edits
                            )
                            st.session_state.generated_text = text
                            st.session_state.generated_title = f"{analysis.get('document_subtype', 'Document')} (Mimicked)"
                            st.toast("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ Document generated!", icon="ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Generation failed: {e}")


def render_guided_workflow():
    """Workflow 2: AI-guided interactive builder."""
    st.markdown("""
    <div style="margin-bottom: 1rem;">
        <h3 style="margin-bottom: 0.5rem;">AI-Guided Document Builder</h3>
        <p style="color: var(--text-secondary);">
            Select a document type and the AI will guide you through all the necessary information
            with contextual questions.
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Initialize session state for guided workflow
    if "guided_step" not in st.session_state:
        st.session_state.guided_step = 0
        st.session_state.guided_answers = {}
        st.session_state.guided_fields = []

    # Step 0: Select document type
    if st.session_state.guided_step == 0:
        doc_type = st.selectbox(
            "What type of document do you want to create?",
            options=list(DOCUMENT_TYPES.keys()),
            format_func=lambda k: DOCUMENT_TYPES[k]["label"],
            key="guided_doc_type"
        )

        doc_def = DOCUMENT_TYPES[doc_type]
        st.markdown(f"*{doc_def['description']}*")

        if st.button("Start Building ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¾Ãƒâ€šÃ‚Â¢", type="primary", use_container_width=True):
            st.session_state.guided_step = 1
            st.session_state.guided_doc_type = doc_type
            st.session_state.guided_fields = doc_def["fields"]
            st.session_state.guided_current_field = 0
            st.rerun()

    # Step 1+: Ask questions one by one
    elif st.session_state.guided_step > 0:
        fields = st.session_state.guided_fields
        current_idx = st.session_state.guided_current_field
        doc_type = st.session_state.guided_doc_type

        if current_idx < len(fields):
            # Show progress
            progress = (current_idx) / len(fields)
            st.progress(progress)
            st.caption(f"Question {current_idx + 1} of {len(fields)}")

            field = fields[current_idx]

            st.markdown(f"### {field['label']}")
            if field.get("help"):
                st.caption(field["help"])

            # Show the appropriate input type
            if field.get("type") == "date":
                answer = st.date_input(
                    "Select date",
                    value=date.today(),
                    key=f"guided_q_{current_idx}",
                    label_visibility="collapsed"
                )
            elif field.get("type") == "textarea":
                answer = st.text_area(
                    "Your answer",
                    placeholder=field.get("placeholder", ""),
                    height=150,
                    key=f"guided_q_{current_idx}",
                    label_visibility="collapsed"
                )
            else:
                answer = st.text_input(
                    "Your answer",
                    placeholder=field.get("placeholder", ""),
                    key=f"guided_q_{current_idx}",
                    label_visibility="collapsed"
                )

            col1, col2 = st.columns([1, 1])

            with col1:
                if current_idx > 0:
                    if st.button("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Previous", use_container_width=True):
                        st.session_state.guided_current_field -= 1
                        st.rerun()

            with col2:
                next_label = "Next ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¾Ãƒâ€šÃ‚Â¢" if current_idx < len(fields) - 1 else "Generate Document ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¨"
                if st.button(next_label, type="primary", use_container_width=True):
                    # Store answer
                    st.session_state.guided_answers[field["key"]] = answer

                    if current_idx < len(fields) - 1:
                        # Move to next question
                        st.session_state.guided_current_field += 1
                        st.rerun()
                    else:
                        # All questions answered - generate document
                        with st.spinner("Generating your document..."):
                            llm = get_llm()
                            collection = get_collection()
                            generator = DocumentGenerator(llm, collection, knowledge_db=get_knowledge_db())

                            try:
                                text = generator.generate(
                                    doc_type,
                                    st.session_state.guided_answers,
                                    use_sec=False,
                                )
                                st.session_state.generated_text = text
                                st.session_state.generated_title = f"{DOCUMENT_TYPES[doc_type]['label']} ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Draft"

                                # Reset guided workflow
                                st.session_state.guided_step = 0
                                st.session_state.guided_answers = {}
                                st.session_state.guided_current_field = 0

                                st.toast("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ Document generated!", icon="ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Generation failed: {e}")

            # Show summary of answers so far
            if st.session_state.guided_answers:
                with st.expander("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¹ Review Your Answers"):
                    for key, value in st.session_state.guided_answers.items():
                        if value:
                            st.caption(f"**{key.replace('_', ' ').title()}:** {value}")


def render_quick_workflow():
    """Workflow 3: Quick static form (original workflow)."""
    st.markdown("""
    <div style="margin-bottom: 1rem;">
        <h3 style="margin-bottom: 0.5rem;">Quick Generate</h3>
        <p style="color: var(--text-secondary);">
            For users who know exactly what they need. Fill out the form and generate immediately.
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Document type selector
    doc_type = st.selectbox(
        "Document type",
        options=list(DOCUMENT_TYPES.keys()),
        format_func=lambda k: DOCUMENT_TYPES[k]["label"],
        key="quick_doc_type",
    )

    doc_def = DOCUMENT_TYPES[doc_type]
    st.markdown(f"*{doc_def['description']}*")

    st.divider()

    # Dynamic form
    params = {}
    with st.form("quick_gen_form"):
        st.markdown("### ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¹ Document Details")
        st.caption("Fill in the information below. The AI will use these details to draft your document.")

        for field in doc_def["fields"]:
            key = field["key"]
            label = field["label"]
            ftype = field.get("type", "text")
            placeholder = field.get("placeholder", "")
            help_text = field.get("help", "")

            if ftype == "date":
                params[key] = st.date_input(
                    label,
                    value=date.today(),
                    key=f"quick_{key}",
                    help=help_text if help_text else None
                )
            elif ftype == "textarea":
                params[key] = st.text_area(
                    label,
                    placeholder=placeholder,
                    key=f"quick_{key}",
                    height=120,
                    help=help_text if help_text else "Provide detailed information here"
                )
            else:
                params[key] = st.text_input(
                    label,
                    placeholder=placeholder,
                    key=f"quick_{key}",
                    help=help_text if help_text else None
                )

        st.divider()
        st.markdown("### ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Optional: Enhance with External Data")
        st.caption("Pull additional context from external databases (optional)")

        use_sec = st.checkbox(
            "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â  Include SEC EDGAR data",
            help="Fetch relevant public company filings from SEC EDGAR database",
            key="quick_use_sec"
        )

        st.divider()
        submitted = st.form_submit_button(
            "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¨ Generate Document",
            type="primary",
            use_container_width=True,
            help="This may take 30-60 seconds depending on the document complexity"
        )

    if submitted:
        llm = get_llm()
        if not llm.is_available():
            st.error(
                "LLM is not available. Check that Ollama is running or OpenAI key is set in Settings."
            )
            return

        collection = get_collection()
        generator = DocumentGenerator(llm, collection, knowledge_db=get_knowledge_db())
        sec_client = get_sec_client() if use_sec else None

        with st.spinner("Generating document... this may take a moment."):
            try:
                text = generator.generate(
                    doc_type,
                    params,
                    sec_client=sec_client,
                    use_sec=use_sec,
                )
                st.session_state.generated_text = text
                st.session_state.generated_title = (
                    f"{doc_def['label']} ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â {params.get('party_a', '') or params.get('entity_name', '') or params.get('case_caption', '') or params.get('re', '') or 'Draft'}"
                )
                st.toast("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ Document generated successfully!", icon="ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦")
            except Exception as e:
                st.error(f"Generation failed: {e}")


# ======================================================================
# TAB 2 ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Generate Document (Main Function)
# ======================================================================

def render_generate_tab():
    """Render the document generation interface with three workflow options."""
    st.markdown("""
    <div class="card-header">
        ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Generate Legal Documents
    </div>
    """, unsafe_allow_html=True)

    # Check LLM availability
    llm = get_llm()
    if not llm.is_available():
        if st.session_state.llm_provider == "ollama":
            st.markdown("""
            <div class="error-card">
                <strong>Ollama Not Running</strong><br><br>
                Start Ollama (<code>ollama serve</code>) and refresh this page, or switch providers in Settings.
            </div>
            """, unsafe_allow_html=True)
        elif st.session_state.llm_provider == "hf_local":
            reason = st.session_state.get("hf_local_error", "HF Local model is not available. Check Settings.")
            st.markdown(f"""
            <div class="error-card">
                <strong>HF Local Model Unavailable</strong><br><br>
                {html_lib.escape(reason)}
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="error-card">
                <strong>OpenAI API Key Required</strong><br><br>
                Add a valid API key in Settings before generating documents.
            </div>
            """, unsafe_allow_html=True)

# ======================================================================
# NEW WORKFLOW: Two-Path Landing Page
# ======================================================================

def render_landing_page():
    """Post-login generation workspace focused on three primary tasks."""
    render_workflow_header(
        "Document Workspace",
        "Select a workflow to continue.",
        step_note="Tip: Start with New Document for drafting, Review Existing for revisions, or Learn to improve retrieval.",
    )

    col1, col2, col3 = st.columns(3, gap="medium")

    with col1:
        st.markdown("""
        <div class="workspace-tile">
            <div class="workspace-tile-title">New Document</div>
            <div class="workspace-tile-copy">Draft a new document with your saved configuration profile.</div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Start Draft", type="secondary", use_container_width=True, key="home_new_doc"):
            st.session_state.workflow_mode = "create"
            st.rerun()

    with col2:
        st.markdown("""
        <div class="workspace-tile">
            <div class="workspace-tile-title">Review Existing</div>
            <div class="workspace-tile-copy">Refine uploaded documents with iterative, in-screen edits.</div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Open Editor", type="secondary", use_container_width=True, key="home_review_doc"):
            st.session_state.workflow_mode = "edit"
            st.rerun()

    with col3:
        st.markdown("""
        <div class="workspace-tile">
            <div class="workspace-tile-title">Learn from Document</div>
            <div class="workspace-tile-copy">Index approved examples to improve future generation quality.</div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Learn From Files", type="secondary", use_container_width=True, key="home_learn_doc"):
            st.session_state.workflow_mode = "learn"
            st.rerun()


# ======================================================================
# PATH A: Edit Existing Document Workflow
# ======================================================================

def render_edit_workflow():
    """Executive edit flow with intent presets and checkpointed revisions."""

    if st.button("Back to Home", key="back_from_edit"):
        st.session_state.workflow_mode = None
        st.session_state.pop("edit_document_text", None)
        st.session_state.pop("edit_document_history", None)
        st.session_state.pop("edit_chat_messages", None)
        st.session_state.pop("edit_revision_goal", None)
        st.session_state.pop("edit_original_file_bytes", None)
        st.session_state.pop("edit_document_original", None)
        st.session_state.pop("edit_edgar_suggestions", None)
        st.session_state.pop("edit_edgar_compare_meta", None)
        st.session_state.pop("edit_edgar_documents", None)
        st.session_state.pop("edit_edgar_active_doc_index", None)
        st.session_state.pop("edit_edgar_search_results", None)
        st.session_state.pop("edit_redline_edit_area", None)
        st.session_state.pop("edit_redline_source_text", None)
        st.session_state.pop("edit_preview_redline_area", None)
        st.session_state.pop("edit_ai_proposal_text", None)
        st.session_state.pop("edit_ai_proposal_base", None)
        st.session_state.pop("edit_ai_proposal_note", None)
        st.rerun()

    if "edit_revision_goal" not in st.session_state:
        st.session_state.edit_revision_goal = "Preserve structure; improve precision and legal clarity."
    if "edit_edgar_suggestions" not in st.session_state:
        st.session_state.edit_edgar_suggestions = []
    if "edit_edgar_compare_meta" not in st.session_state:
        st.session_state.edit_edgar_compare_meta = None
    if "edit_edgar_documents" not in st.session_state:
        st.session_state.edit_edgar_documents = []
    if "edit_edgar_active_doc_index" not in st.session_state:
        st.session_state.edit_edgar_active_doc_index = 0
    if "edit_edgar_search_results" not in st.session_state:
        st.session_state.edit_edgar_search_results = []
    if "edit_redline_source_text" not in st.session_state:
        st.session_state.edit_redline_source_text = ""
    if "edit_ai_proposal_text" not in st.session_state:
        st.session_state.edit_ai_proposal_text = None
    if "edit_ai_proposal_base" not in st.session_state:
        st.session_state.edit_ai_proposal_base = ""
    if "edit_ai_proposal_note" not in st.session_state:
        st.session_state.edit_ai_proposal_note = ""

    if "edit_document_text" not in st.session_state:
        render_workflow_header("Edit Existing Document", "Upload one document and apply targeted revisions.", progress=0.33, step_note="Step 1 of 3: Upload")
        st.markdown("""
        <div class="info-card">
            Upload one document or paste a Google Docs URL to begin. The assistant maintains full-document continuity for each revision.
        </div>
        """, unsafe_allow_html=True)

        uploaded_file = st.file_uploader(
            "Upload document",
            type=["pdf", "doc", "docx", "docm", "txt"],
            key="edit_upload",
            help="Supported: PDF, DOC, DOCX, DOCM, TXT"
        )

        st.markdown("#### Or import from Google Docs")
        google_doc_url = st.text_input(
            "Google Doc URL",
            placeholder="https://docs.google.com/document/d/<DOC_ID>/edit",
            key="edit_google_doc_url",
            help="The Google Doc should be accessible to import (for local testing, use 'Anyone with the link').",
        )
        import_google_doc = st.button(
            "Import Google Doc",
            key="edit_import_google_doc",
            use_container_width=True,
        )

        loaded_text = None
        loaded_name = None
        loaded_ext = ".txt"
        loaded_original_bytes = None

        if uploaded_file:
            raw_bytes = uploaded_file.getvalue()
            with st.spinner("Reading document..."):
                text = extract_text_from_file(io.BytesIO(raw_bytes), uploaded_file.name)

            if text.startswith("Error reading file:"):
                st.error(text)
            elif len(text) < 50:
                st.error("Document appears too short or unreadable. Upload a valid file.")
            else:
                ext = ("." + uploaded_file.name.split(".")[-1].lower()) if "." in uploaded_file.name else ".txt"
                loaded_text = text
                loaded_name = uploaded_file.name
                loaded_ext = ext
                loaded_original_bytes = raw_bytes if ext in [".docx", ".docm"] else None

        if import_google_doc:
            if not google_doc_url.strip():
                st.error("Enter a Google Docs URL to import.")
            else:
                with st.spinner("Importing Google Doc..."):
                    text = extract_text_from_google_doc_url(google_doc_url.strip())
                if text.startswith("Error reading file:"):
                    st.error(text)
                elif len(text) < 50:
                    st.error("Imported Google Doc appears too short or unreadable.")
                else:
                    doc_id = _extract_google_doc_id(google_doc_url.strip())
                    loaded_text = text
                    loaded_name = f"google-doc-{doc_id}.txt" if doc_id else "google-doc.txt"
                    loaded_ext = ".txt"
                    loaded_original_bytes = None

        if loaded_text:
            st.session_state.edit_document_text = loaded_text
            st.session_state.edit_document_original = loaded_text
            st.session_state.edit_document_history = [{"version": 0, "text": loaded_text, "change": "Original document"}]
            st.session_state.edit_chat_messages = []
            st.session_state.edit_filename = loaded_name or "Imported document"
            st.session_state.edit_file_ext = loaded_ext
            st.session_state.edit_original_file_bytes = loaded_original_bytes
            st.session_state.edit_edgar_suggestions = []
            st.session_state.edit_edgar_compare_meta = None
            st.session_state.edit_edgar_documents = []
            st.session_state.edit_edgar_active_doc_index = 0
            st.session_state.edit_edgar_search_results = []
            st.session_state.edit_redline_edit_area = build_diff_edit_markup(loaded_text, loaded_text)[0]
            st.session_state.edit_redline_source_text = loaded_text
            st.session_state.edit_preview_redline_area = st.session_state.edit_redline_edit_area
            st.session_state.edit_ai_proposal_text = None
            st.session_state.edit_ai_proposal_base = ""
            st.session_state.edit_ai_proposal_note = ""
            st.success(f"Loaded {st.session_state.edit_filename} ({len(loaded_text):,} characters)")
            st.rerun()
        return

    render_workflow_header("Edit Existing Document", "Refine your working draft and apply AI revisions.", progress=0.66, step_note="Step 2 of 3: Objective and revisions")
    st.markdown(f"""
    <div class="info-card">
        <strong>Document:</strong> {st.session_state.get('edit_filename', 'Untitled')}<br>
        <strong>Version:</strong> {len(st.session_state.edit_document_history)}<br>
        <strong>Length:</strong> {len(st.session_state.edit_document_text)} characters
    </div>
    """, unsafe_allow_html=True)

    preset_col, objective_col = st.columns([2, 3], gap="large")
    with preset_col:
        st.markdown("#### Revision Preset")
        preset = st.radio(
            "Choose preset",
            options=["Tighten", "Formalize", "Clarify", "Shorten", "Custom"],
            index=4,
            key="edit_preset_radio",
            label_visibility="collapsed"
        )
        if preset == "Tighten":
            st.session_state.edit_revision_goal = "Reduce verbosity and tighten clause language without changing legal meaning."
        elif preset == "Formalize":
            st.session_state.edit_revision_goal = "Use formal legal register and improve professional tone."
        elif preset == "Clarify":
            st.session_state.edit_revision_goal = "Clarify ambiguous provisions and improve definitional precision."
        elif preset == "Shorten":
            st.session_state.edit_revision_goal = "Shorten text where possible while preserving obligations and risk protections."

    with objective_col:
        st.markdown("#### Current Objective")
        st.session_state.edit_revision_goal = st.text_area(
            "Revision objective",
            value=st.session_state.edit_revision_goal,
            height=120,
            label_visibility="collapsed",
            key="edit_goal_area",
            help="This objective is applied to each AI revision request."
        )
        if len(st.session_state.edit_revision_goal.strip()) < 12:
            st.caption("Add more detail for higher-quality revisions.")

    col_preview, col_ai = st.columns([3, 2], gap="large")

    with col_preview:
        st.markdown("### Working Draft")
        edited_text = st.text_area(
            "Working draft",
            value=st.session_state.edit_document_text,
            height=560,
            key="edit_preview_area",
            label_visibility="collapsed"
        )
        if edited_text != st.session_state.edit_document_text:
            st.session_state.edit_document_text = edited_text
            st.session_state.edit_redline_source_text = edited_text
            st.session_state.edit_redline_edit_area = build_diff_edit_markup(
                st.session_state.get("edit_document_original", ""),
                edited_text,
            )[0]
            st.session_state.edit_preview_redline_area = st.session_state.edit_redline_edit_area

        edgar_docs = st.session_state.get("edit_edgar_documents", [])
        if edgar_docs:
            st.markdown("#### EDGAR Reference Document")
            active_idx = st.selectbox(
                "Retrieved EDGAR document",
                options=list(range(len(edgar_docs))),
                index=min(st.session_state.get("edit_edgar_active_doc_index", 0), len(edgar_docs) - 1),
                format_func=lambda i: (
                    f"{i + 1}. {edgar_docs[i].get('entity_name', 'Unknown')} "
                    f"| {edgar_docs[i].get('form_type', 'N/A')} | {edgar_docs[i].get('file_date', 'N/A')}"
                ),
                key="edit_edgar_doc_selector",
            )
            st.session_state.edit_edgar_active_doc_index = active_idx
            active_doc = edgar_docs[active_idx]
            if active_doc.get("url"):
                st.markdown(f"[Open SEC filing source]({active_doc['url']})")
            st.text_area(
                "EDGAR document text",
                value=active_doc.get("text", ""),
                height=220,
                disabled=True,
                key=f"edit_edgar_doc_text_{active_idx}",
            )

        proposal_text = st.session_state.get("edit_ai_proposal_text")
        proposal_base = st.session_state.get("edit_ai_proposal_base", "")
        if proposal_text:
            st.markdown("#### AI Proposed Changes (Pending Commit)")
            proposal_html, proposal_add, proposal_del = build_diff_highlight_html(
                proposal_base,
                proposal_text,
            )
            st.markdown(
                f"<div style='white-space:pre-wrap;border:1px solid #d7dbd7;border-radius:10px;padding:0.9rem;background:#ffffff;line-height:1.55;'>{proposal_html}</div>",
                unsafe_allow_html=True,
            )
            st.caption(f"Proposed delta (+{proposal_add} / -{proposal_del})")
            pc1, pc2 = st.columns(2)
            with pc1:
                if st.button("Commit Proposed Changes", key="edit_commit_ai_proposal", use_container_width=True):
                    version_num = len(st.session_state.edit_document_history)
                    st.session_state.edit_document_text = proposal_text
                    st.session_state.edit_document_history.append({
                        "version": version_num,
                        "text": proposal_text,
                        "change": st.session_state.get("edit_ai_proposal_note", "AI proposed revision")[:72],
                        "kind": "ai",
                    })
                    st.session_state.edit_ai_proposal_text = None
                    st.session_state.edit_ai_proposal_base = ""
                    st.session_state.edit_ai_proposal_note = ""
                    st.rerun()
            with pc2:
                if st.button("Discard Proposed Changes", key="edit_discard_ai_proposal", use_container_width=True):
                    st.session_state.edit_ai_proposal_text = None
                    st.session_state.edit_ai_proposal_base = ""
                    st.session_state.edit_ai_proposal_note = ""
                    st.rerun()

        diff_html, add_count, del_count = build_diff_highlight_html(
            st.session_state.get("edit_document_original", ""),
            st.session_state.edit_document_text,
        )

        history = st.session_state.edit_document_history
        ai_entries = [
            v for v in history
            if v.get("kind") == "ai" or (
                v.get("version", 0) > 0
                and "imported" not in str(v.get("change", "")).lower()
                and "baseline" not in str(v.get("change", "")).lower()
                and "original" not in str(v.get("change", "")).lower()
            )
        ]
        if ai_entries:
            latest_ai_text = ai_entries[-1]["text"]
            latest_ai_base = ai_entries[-2]["text"] if len(ai_entries) > 1 else st.session_state.get("edit_document_original", "")
            latest_diff_html, latest_add_count, latest_del_count = build_diff_highlight_html(
                latest_ai_base,
                latest_ai_text,
            )
            st.markdown(f"#### Latest AI Revision Commit Delta (+{latest_add_count} / -{latest_del_count})")
            st.markdown(
                f"<div style='white-space:pre-wrap;border:1px solid #d7dbd7;border-radius:10px;padding:0.9rem;background:#ffffff;line-height:1.55;'>{latest_diff_html}</div>",
                unsafe_allow_html=True,
            )

        st.markdown(f"#### Total Revisions vs Original (+{add_count} / -{del_count})")
        st.markdown(
            f"<div style='white-space:pre-wrap;border:1px solid #d7dbd7;border-radius:10px;padding:0.9rem;background:#ffffff;line-height:1.55;'>{diff_html}</div>",
            unsafe_allow_html=True,
        )

        a, b, c = st.columns(3)
        with a:
            if len(st.session_state.edit_document_history) > 1 and st.button("Undo", use_container_width=True, key="edit_undo"):
                st.session_state.edit_document_history.pop()
                st.session_state.edit_document_text = st.session_state.edit_document_history[-1]["text"]
                st.rerun()
        with b:
            original_name = st.session_state.get("edit_filename", "document")
            ext = st.session_state.get("edit_file_ext", ".txt")
            base_name = original_name.rsplit(".", 1)[0] if "." in original_name else original_name

            if ext == ".txt":
                st.download_button(
                    label="Download clean final (.txt)",
                    data=st.session_state.edit_document_text.encode("utf-8"),
                    file_name=f"{base_name}_edited.txt",
                    mime="text/plain",
                    use_container_width=True,
                )
            elif ext in [".docx", ".docm"] and st.session_state.get("edit_original_file_bytes"):
                try:
                    preserved_bytes = DocumentGenerator.replace_text_preserve_word_template(
                        st.session_state.edit_original_file_bytes,
                        st.session_state.edit_document_text,
                    )
                    out_ext = ext
                    out_mime = (
                        "application/vnd.ms-word.document.macroEnabled.12"
                        if out_ext == ".docm"
                        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.download_button(
                        label=f"Download clean final (preserved {out_ext})",
                        data=preserved_bytes,
                        file_name=f"{base_name}_edited{out_ext}",
                        mime=out_mime,
                        use_container_width=True,
                    )
                except Exception as preserve_err:
                    st.warning(f"Could not preserve original package exactly ({preserve_err}). Using standard DOCX export.")
                    docx_bytes = DocumentGenerator.text_to_docx(
                        st.session_state.edit_document_text,
                        original_name,
                    )
                    st.download_button(
                        label="Download clean final (.docx)",
                        data=docx_bytes,
                        file_name=f"{base_name}_edited.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
            else:
                docx_bytes = DocumentGenerator.text_to_docx(
                    st.session_state.edit_document_text,
                    original_name
                )
                st.download_button(
                    label="Download clean final (.docx)",
                    data=docx_bytes,
                    file_name=f"{base_name}_edited.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                if ext == ".pdf":
                    st.caption("Original was PDF. Export is provided as editable DOCX.")
                elif ext == ".doc":
                    st.caption("Original was legacy .doc. Clean and redline exports are provided as .docx.")

            redline_bytes = DocumentGenerator.diff_to_docx(
                st.session_state.get("edit_document_original", ""),
                st.session_state.edit_document_text,
                title=f"{base_name} Redline",
            )
            st.download_button(
                label="Download redline (.docx)",
                data=redline_bytes,
                file_name=f"{base_name}_redline.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="edit_redline_download",
            )
        with c:
            if st.button("Reset", use_container_width=True, key="edit_reset"):
                st.session_state.pop("edit_document_text", None)
                st.session_state.pop("edit_document_history", None)
                st.session_state.pop("edit_chat_messages", None)
                st.session_state.pop("edit_original_file_bytes", None)
                st.session_state.pop("edit_document_original", None)
                st.session_state.pop("edit_edgar_suggestions", None)
                st.session_state.pop("edit_edgar_compare_meta", None)
                st.session_state.pop("edit_edgar_documents", None)
                st.session_state.pop("edit_edgar_active_doc_index", None)
                st.session_state.pop("edit_edgar_search_results", None)
                st.session_state.pop("edit_redline_edit_area", None)
                st.session_state.pop("edit_redline_source_text", None)
                st.session_state.pop("edit_preview_redline_area", None)
                st.session_state.pop("edit_ai_proposal_text", None)
                st.session_state.pop("edit_ai_proposal_base", None)
                st.session_state.pop("edit_ai_proposal_note", None)
                st.rerun()

        if len(st.session_state.edit_document_history) > 1:
            with st.expander("Revision checkpoints"):
                for version in reversed(st.session_state.edit_document_history):
                    st.caption(f"v{version['version']}: {version['change']}")

    with col_ai:
        st.markdown("### AI Revision Assistant")
        st.caption("Use specific instructions. The assistant returns a full revised document each time.")

        with st.expander("Compare to SEC EDGAR", expanded=False):
            sec_client = get_sec_client()
            default_query = st.session_state.get("edit_filename", "")
            if "." in default_query:
                default_query = default_query.rsplit(".", 1)[0]

            edgar_query = st.text_input(
                "EDGAR search query",
                value=st.session_state.get("edit_edgar_query", default_query),
                key="edit_edgar_query",
                help="Search for comparable filings (example: software services agreement or public company name).",
            )
            edgar_forms_raw = st.text_input(
                "Form types (comma-separated, optional)",
                value=st.session_state.get("edit_edgar_forms", "8-K,10-K,10-Q"),
                key="edit_edgar_forms",
            )
            max_results = st.slider("Top comparable results", min_value=1, max_value=10, value=10, key="edit_edgar_max_results")

            if not sec_client.is_configured():
                st.info("SEC EDGAR User-Agent is not configured. Add it in Settings to enable comparisons.")

            if st.button("Search Comparable EDGAR Documents", use_container_width=True, key="edit_run_edgar_compare"):
                if not sec_client.is_configured():
                    st.error("Set SEC EDGAR User-Agent in Settings first.")
                else:
                    form_types = [f.strip().upper() for f in edgar_forms_raw.split(",") if f.strip()]
                    with st.spinner("Searching comparable EDGAR filings..."):
                        result = search_edgar_comparable_documents(
                            sec_client=sec_client,
                            query=edgar_query,
                            form_types=form_types,
                            max_results=max_results,
                        )
                    if result.get("error"):
                        st.error(result["error"])
                        st.session_state.edit_edgar_search_results = []
                        st.session_state.edit_edgar_compare_meta = None
                    else:
                        st.session_state.edit_edgar_search_results = result.get("results", [])
                        st.session_state.edit_edgar_compare_meta = {
                            "hits_considered": result.get("hits_considered", 0),
                            "hits_listed": result.get("hits_listed", 0),
                        }
                        st.success(f"Listed {len(st.session_state.edit_edgar_search_results)} comparable EDGAR result(s).")

            compare_meta = st.session_state.get("edit_edgar_compare_meta")
            if compare_meta:
                st.caption(
                    f"Top results listed: {compare_meta.get('hits_listed', 0)} / "
                    f"{compare_meta.get('hits_considered', 0)} matched"
                )

            search_results = st.session_state.get("edit_edgar_search_results", [])
            if search_results:
                st.markdown("#### Comparable EDGAR Results")
                for i, filing in enumerate(search_results):
                    c1, c2 = st.columns([4, 1])
                    with c1:
                        st.markdown(
                            f"**{i + 1}. {filing.get('entity_name', 'Unknown')}** | "
                            f"{filing.get('form_type', 'N/A')} | {filing.get('file_date', 'N/A')}"
                        )
                        if filing.get("url"):
                            st.markdown(f"[Open filing index]({filing['url']})")
                    with c2:
                        if st.button("Load Preview", key=f"edit_edgar_load_{i}", use_container_width=True):
                            with st.spinner("Loading EDGAR filing preview..."):
                                loaded = load_edgar_document_preview(sec_client, filing)
                            if not loaded:
                                st.warning("Could not load usable text for that filing.")
                            else:
                                existing_urls = {d.get("url", "") for d in st.session_state.get("edit_edgar_documents", [])}
                                if loaded.get("url") not in existing_urls:
                                    st.session_state.edit_edgar_documents.append(loaded)
                                    st.session_state.edit_edgar_active_doc_index = len(st.session_state.edit_edgar_documents) - 1
                                else:
                                    for idx, d in enumerate(st.session_state.edit_edgar_documents):
                                        if d.get("url") == loaded.get("url"):
                                            st.session_state.edit_edgar_active_doc_index = idx
                                            break
                                st.success("Loaded into EDGAR Reference Document preview.")
                                st.rerun()

                st.caption("Use the Working Draft 'EDGAR Reference Document' picker to preview loaded filings for copy/paste drafting.")

        for msg in st.session_state.edit_chat_messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        if st.button("Clear revision chat", use_container_width=True, key="edit_clear_chat"):
            st.session_state.edit_chat_messages = []
            st.rerun()

        if prompt := st.chat_input("Example: Clarify notice mechanics and align defined terms"):
            st.session_state.edit_chat_messages.append({"role": "user", "content": prompt})

        if st.session_state.edit_chat_messages and st.session_state.edit_chat_messages[-1]["role"] == "user":
            pending_prompt = st.session_state.edit_chat_messages[-1]["content"]
            with st.chat_message("assistant"):
                with st.spinner("Applying revision..."):
                    llm = get_llm()
                    messages = [
                        {
                            "role": "system",
                            "content": (
                                "You are a legal document editing assistant. Return the FULL revised document. "
                                "Preserve enforceability and structure unless explicitly asked to alter them."
                            )
                        },
                        {
                            "role": "user",
                            "content": (
                                f"Revision objective: {st.session_state.edit_revision_goal}\n\n"
                                f"Current document:\n\n{st.session_state.edit_document_text}\n\n"
                                f"Change request: {pending_prompt}\n\n"
                                "Return only the full updated document text."
                            )
                        }
                    ]
                    try:
                        proposal_input_text = st.session_state.get("edit_ai_proposal_text") or st.session_state.edit_document_text
                        proposal_base_text = st.session_state.get("edit_ai_proposal_base") or st.session_state.edit_document_text
                        messages[1]["content"] = (
                            f"Revision objective: {st.session_state.edit_revision_goal}\n\n"
                            f"Current document:\n\n{proposal_input_text}\n\n"
                            f"Change request: {pending_prompt}\n\n"
                            "Return only the full updated document text."
                        )
                        response = llm.chat(messages, temperature=0.2, max_tokens=4096)
                        st.session_state.edit_ai_proposal_text = response
                        st.session_state.edit_ai_proposal_base = proposal_base_text
                        st.session_state.edit_ai_proposal_note = pending_prompt
                        st.session_state.edit_chat_messages.append({"role": "assistant", "content": "Proposed revision ready. Review in 'AI Proposed Changes' and commit or request clarification."})
                        st.rerun()
                    except Exception as e:
                        err = f"Error: {e}"
                        st.error(err)
                        st.session_state.edit_chat_messages.append({"role": "assistant", "content": err})


# ======================================================================
# PATH B: Create New Document Workflow
# ======================================================================

def render_create_workflow():
    """Create flow with short setup wizard + guided drafting workspace."""

    if st.button("Back to Home", key="back_from_create"):
        st.session_state.workflow_mode = None
        for k in [
            "create_setup_complete", "create_setup_goal", "create_setup_style", "create_setup_guidance", "create_setup_style_example", "create_setup_style_doc_id", "create_setup_style_doc_name",
            "create_doc_type", "create_doc_type_label", "create_fields", "create_generated_text", "create_chat_messages",
            "create_doc_description", "create_detection_reason", "create_doc_description_input",
            "create_original_generated_text", "create_redline_edit_area"
        ]:
            st.session_state.pop(k, None)
        st.rerun()

    def _preview_value(value):
        if value is None:
            return ""
        return str(value).strip()

    def _build_live_preview(doc_label: str, payload: dict) -> str:
        goal = _preview_value(st.session_state.get("create_setup_goal"))
        style_source = _preview_value(st.session_state.get("create_setup_style"))
        style_example = _preview_value(st.session_state.get("create_setup_style_example"))
        guidance = _preview_value(st.session_state.get("create_setup_guidance"))
        non_empty_items = [(k, _preview_value(v)) for k, v in payload.items() if _preview_value(v)]

        lines = [
            doc_label.upper(),
            "",
            f"Draft objective: {goal or 'Not set'}",
            f"Style source: {style_source or 'Not set'}",
            f"Style example: {'Provided' if style_example else 'Not provided'}",
            f"Guidance mode: {guidance or 'Not set'}",
            "",
            "Key terms:"
        ]
        if non_empty_items:
            for key, value in non_empty_items[:8]:
                pretty = key.replace("_", " ").title()
                lines.append(f"- {pretty}: {value}")
        else:
            lines.append("- Waiting for input fields")

        lines.extend([
            "",
            "Preview clause:",
            "This draft is generated from the provided inputs and is presented for attorney review."
        ])
        return "\n".join(lines)

    if not st.session_state.get("create_setup_complete", False):
        render_workflow_header("Create New Document", "Configure drafting profile and generate a first pass.", progress=0.33, step_note="Step 1 of 3: Draft profile")
        left, right = st.columns([3, 2], gap="large")
        with left:
            st.markdown("### Step 1: Draft Profile")
            st.session_state.create_setup_goal = st.selectbox(
                "Drafting goal",
                ["Board-ready first draft", "Client-ready polished draft", "Fast internal draft"],
                index=0,
                key="create_setup_goal_select",
                help="Controls drafting strictness and polish level."
            )
            st.session_state.create_setup_style = st.selectbox(
                "Style source",
                [
                    "Learned templates first",
                    "Standard template library",
                    "Custom freeform",
                    "User-provided style example"
                ],
                index=0,
                key="create_setup_style_select",
                help="Controls how much the assistant relies on known template patterns."
            )
            requires_style_example = st.session_state.create_setup_style == "User-provided style example"
            style_example_help = (
                "Paste representative language. The assistant will prioritize this style when drafting."
                if requires_style_example
                else "Optional. Add representative language if you want the draft tone and structure to mirror it."
            )
            st.session_state.create_setup_style_example = st.text_area(
                "Style example text",
                value=st.session_state.get("create_setup_style_example", ""),
                placeholder="Paste a short sample clause or paragraph to emulate tone and structure.",
                height=140,
                key="create_setup_style_example_input",
                help=style_example_help,
            )
            st.session_state.create_setup_guidance = st.selectbox(
                "Guidance level",
                ["Inline guidance", "Minimal guidance"],
                index=0,
                key="create_setup_guidance_select"
            )

            if st.button("Continue to Document Type", type="secondary", use_container_width=True, key="create_setup_continue"):
                if requires_style_example and len(st.session_state.get("create_setup_style_example", "").strip()) < 40:
                    st.warning("Please provide a longer style example so the model can follow it closely.")
                else:
                    st.session_state.create_setup_complete = True
                    st.rerun()
        with right:
            st.markdown("### Live Preview")
            st.code(_build_live_preview("Document Draft", {}), language="markdown")
        return

    if "create_doc_type" not in st.session_state:
        render_workflow_header("Create New Document", "Describe your request so the app can infer the right document type.", progress=0.66, step_note="Step 2 of 3: Document type")
        left, right = st.columns([3, 2], gap="large")

        with left:
            st.markdown("### Step 2: Describe The Document")
            description = st.text_area(
                "Describe what you need",
                value=st.session_state.get("create_doc_description_input", ""),
                placeholder="Example: Create a mutual NDA between two SaaS companies with a 3-year term, Delaware law, and carve-outs for independently developed information.",
                height=180,
                key="create_doc_description_input",
                help="Use plain language. The system will infer the document type and generate the right fields.",
            )

            if st.button("Analyze Description", type="secondary", use_container_width=True, key="create_doc_type_confirm"):
                if len(description.strip()) < 24:
                    st.warning("Please provide a fuller description so the model can infer the right document type.")
                else:
                    selected, reason = infer_document_type_from_description(description)
                    st.session_state.create_doc_description = description.strip()
                    st.session_state.create_detection_reason = reason

                    if selected == "custom_document":
                        st.session_state.create_doc_type = selected
                        st.session_state.create_doc_type_label = "Custom Document"
                        with st.spinner("Analyzing description and building fields..."):
                            llm = get_llm()
                            generator = DocumentGenerator(llm, get_collection(), knowledge_db=get_knowledge_db())
                            fields = generator.get_required_fields_for_type("Custom Document")
                            st.session_state.create_fields = fields
                            st.session_state.create_is_learned = any(f.get("learned") for f in fields)
                    else:
                        st.session_state.create_doc_type = selected
                        st.session_state.create_doc_type_label = DOCUMENT_TYPES[selected]["label"]
                        st.session_state.create_fields = DOCUMENT_TYPES[selected]["fields"]
                        st.session_state.create_is_learned = False

                    st.success(f"Detected: {st.session_state.create_doc_type_label}")
                    st.rerun()

        with right:
            detected_label = st.session_state.get("create_doc_type_label", "Pending detection")
            st.markdown("### Live Preview")
            st.code(_build_live_preview(detected_label if detected_label != "Pending detection" else "Document Draft", {
                "description": st.session_state.get("create_doc_description_input", "")
            }), language="markdown")
            if st.session_state.get("create_detection_reason"):
                st.caption(f"Detection: {st.session_state.get('create_detection_reason')}")
        return

    if "create_generated_text" not in st.session_state:
        render_workflow_header("Create New Document", "Fill key facts and generate your first draft.", progress=1.0, step_note="Step 3 of 3: Generate")
        st.markdown(f"""
        <div class="info-card">
            <strong>Step 3 of 3:</strong> Enter facts and watch the draft preview update live.<br>
            <strong>Profile:</strong> {st.session_state.get('create_setup_goal')} | {st.session_state.get('create_setup_style')} | {st.session_state.get('create_setup_guidance')}<br>
            <strong>Style example:</strong> {"Provided" if st.session_state.get('create_setup_style_example', '').strip() else "Not provided"}<br>`r`n            <strong>Style document:</strong> {st.session_state.get('create_setup_style_doc_name', "") or "Not selected"}
        </div>
        """, unsafe_allow_html=True)

        left, right = st.columns([3, 2], gap="large")
        form_data = {}
        fields = st.session_state.create_fields

        with left:
            st.caption(f"Fields loaded: {len(fields)}")
            for field in fields:
                key = field.get("key", "field")
                label = field.get("label", key.replace("_", " ").title())
                field_type = field.get("type", "text")
                placeholder = field.get("placeholder", "")
                help_text = field.get("help", "")
                widget_key = f"create_live_{key}"

                if field_type == "date":
                    form_data[key] = st.date_input(label, value=st.session_state.get(widget_key, date.today()), key=widget_key, help=help_text)
                elif field_type == "textarea":
                    form_data[key] = st.text_area(label, value=st.session_state.get(widget_key, ""), placeholder=placeholder, height=96, key=widget_key, help=help_text)
                else:
                    form_data[key] = st.text_input(label, value=st.session_state.get(widget_key, ""), placeholder=placeholder, key=widget_key, help=help_text)

            llm = get_llm()
            generate_disabled = not llm.is_available()
            if generate_disabled:
                st.caption("Configure an AI provider in Settings to enable draft generation.")
            if st.button("Generate Draft", type="primary", use_container_width=True, key="create_generate_live", disabled=generate_disabled):
                if not llm.is_available():
                    st.error("LLM is not available. Open Settings and configure a provider.")
                else:
                    with st.spinner("Generating draft..."):
                        try:
                            collection = get_collection()
                            generator = DocumentGenerator(llm, collection, knowledge_db=get_knowledge_db())
                            doc_type = st.session_state.create_doc_type

                            generation_context = {
                                "draft_goal": st.session_state.get("create_setup_goal", ""),
                                "style_source": st.session_state.get("create_setup_style", ""),
                                "style_example_text": st.session_state.get("create_setup_style_example", ""),
                                "guidance_level": st.session_state.get("create_setup_guidance", ""),
                                "document_description": st.session_state.get("create_doc_description", ""),
                            }
                            form_payload = {**form_data, **generation_context}

                            if doc_type in DOCUMENT_TYPES:
                                text_out = generator.generate(doc_type, form_payload, use_sec=False)
                            else:
                                system_prompt = (
                                    "You are a legal document drafting assistant. Generate a professional legal draft "
                                    "using formal, structured language. This draft is for attorney review, not legal advice."
                                )
                                params = "\n".join([f"{k}: {v}" for k, v in form_payload.items() if v])
                                user_prompt = f"Generate a {st.session_state.create_doc_type_label} with:\n\n{params}"
                                text_out = llm.generate_document(system_prompt, user_prompt)

                            st.session_state.create_generated_text = text_out
                            st.session_state.create_original_generated_text = text_out
                            st.session_state.create_redline_edit_area = build_diff_edit_markup(text_out, text_out)[0]
                            st.session_state.create_chat_messages = []
                            st.rerun()
                        except Exception as e:
                            st.error(f"Generation failed: {e}")

        with right:
            st.markdown("### Live Preview")
            st.code(_build_live_preview(st.session_state.get("create_doc_type_label", "Document Draft"), form_data), language="markdown")
        return

    st.markdown(f"""
    <div class="info-card">
        <strong>Draft:</strong> {st.session_state.create_doc_type_label}<br>
        <strong>Objective:</strong> {st.session_state.get('create_setup_goal')}
    </div>
    """, unsafe_allow_html=True)

    if "create_original_generated_text" not in st.session_state:
        st.session_state.create_original_generated_text = st.session_state.create_generated_text

    col_preview, col_ai = st.columns([3, 2], gap="large")
    with col_preview:
        st.markdown("### Draft Workspace")
        edited_text = st.text_area(
            "Draft",
            value=st.session_state.create_generated_text,
            height=560,
            key="create_preview_area",
            label_visibility="collapsed"
        )
        if edited_text != st.session_state.create_generated_text:
            st.session_state.create_generated_text = edited_text

        create_diff_html, create_add_count, create_del_count = build_diff_highlight_html(
            st.session_state.get("create_original_generated_text", ""),
            st.session_state.create_generated_text,
        )
        st.markdown(f"#### Redline Preview (+{create_add_count} / -{create_del_count})")
        st.markdown(
            f"<div style='white-space:pre-wrap;border:1px solid #d7dbd7;border-radius:10px;padding:0.9rem;background:#ffffff;line-height:1.55;'>{create_diff_html}</div>",
            unsafe_allow_html=True,
        )

        create_edit_markup, create_markup_add, create_markup_del = build_diff_edit_markup(
            st.session_state.get("create_original_generated_text", ""),
            st.session_state.create_generated_text,
        )
        if "create_redline_edit_area" not in st.session_state:
            st.session_state.create_redline_edit_area = create_edit_markup

        with st.expander(f"Editable Redline Mode (+{create_markup_add} / -{create_markup_del})", expanded=False):
            st.caption("Edit inline using [[+added text+]] and [[-deleted text-]] markers, then apply.")
            cr1, cr2 = st.columns(2)
            with cr1:
                if st.button("Refresh from Current Redline", key="create_redline_refresh", use_container_width=True):
                    st.session_state.create_redline_edit_area = create_edit_markup
                    st.rerun()
            with cr2:
                if st.button("Apply Redline Edits", key="create_redline_apply", use_container_width=True):
                    try:
                        create_new_text = apply_edit_markup_to_text(st.session_state.get("create_redline_edit_area", ""))
                        if create_new_text != st.session_state.create_generated_text:
                            st.session_state.create_generated_text = create_new_text
                            st.session_state.create_preview_area = create_new_text
                            st.session_state.create_redline_edit_area = build_diff_edit_markup(
                                st.session_state.get("create_original_generated_text", ""),
                                create_new_text,
                            )[0]
                            st.rerun()
                        else:
                            st.info("No text changes detected from redline markup.")
                    except Exception as create_redline_err:
                        st.error(f"Could not apply redline edits: {create_redline_err}")

            st.text_area(
                "Editable redline",
                key="create_redline_edit_area",
                height=220,
                label_visibility="collapsed",
            )

        a, b, c = st.columns(3)
        with a:
            if st.button("Start Over", use_container_width=True, key="create_start_over_exec"):
                for k in [
                    "create_setup_style_example",
                    "create_doc_type", "create_fields", "create_generated_text", "create_chat_messages",
                    "create_doc_description", "create_detection_reason", "create_doc_description_input",
                    "create_original_generated_text", "create_redline_edit_area"
                ]:
                    st.session_state.pop(k, None)
                st.rerun()
        with b:
            docx_bytes = DocumentGenerator.text_to_docx(
                st.session_state.create_generated_text,
                st.session_state.create_doc_type_label
            )
            st.download_button(
                label="Download clean final (.docx)",
                data=docx_bytes,
                file_name=f"{st.session_state.create_doc_type_label.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            create_redline_bytes = DocumentGenerator.diff_to_docx(
                st.session_state.get("create_original_generated_text", ""),
                st.session_state.create_generated_text,
                title=f"{st.session_state.create_doc_type_label} Redline",
            )
            st.download_button(
                label="Download redline (.docx)",
                data=create_redline_bytes,
                file_name=f"{st.session_state.create_doc_type_label.replace(' ', '_')}_redline.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="create_redline_download",
            )
        with c:
            if st.button("Open in Editor", use_container_width=True, key="create_open_in_editor"):
                doc_label = st.session_state.get("create_doc_type_label", "Generated_Document")
                baseline_text = st.session_state.get("create_original_generated_text", st.session_state.create_generated_text)
                current_text = st.session_state.create_generated_text
                generated_name = f"{doc_label.replace(' ', '_')}.docx"

                st.session_state.edit_document_original = baseline_text
                st.session_state.edit_document_text = current_text
                if current_text != baseline_text:
                    st.session_state.edit_document_history = [
                        {"version": 0, "text": baseline_text, "change": "Generated draft baseline"},
                        {"version": 1, "text": current_text, "change": "Imported from New Document"},
                    ]
                else:
                    st.session_state.edit_document_history = [
                        {"version": 0, "text": baseline_text, "change": "Generated draft baseline"},
                    ]
                st.session_state.edit_chat_messages = []
                st.session_state.edit_filename = generated_name
                st.session_state.edit_file_ext = ".docx"
                st.session_state.edit_original_file_bytes = DocumentGenerator.text_to_docx(baseline_text, generated_name)
                st.session_state.edit_edgar_suggestions = []
                st.session_state.edit_edgar_compare_meta = None
                st.session_state.edit_edgar_documents = []
                st.session_state.edit_edgar_active_doc_index = 0
                st.session_state.edit_edgar_search_results = []
                st.session_state.edit_ai_proposal_text = None
                st.session_state.edit_ai_proposal_base = ""
                st.session_state.edit_ai_proposal_note = ""
                st.session_state.workflow_mode = "edit"
                st.rerun()

            if st.button("Re-run Guidance", use_container_width=True, key="create_rerun_setup"):
                st.session_state.create_setup_complete = False
                st.rerun()

    with col_ai:
        st.markdown("### AI Revision Assistant")
        st.caption("Request targeted improvements. The assistant returns a full updated draft.")

        for msg in st.session_state.get("create_chat_messages", []):
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        if st.button("Clear revision chat", use_container_width=True, key="create_clear_chat"):
            st.session_state.create_chat_messages = []
            st.rerun()

        if prompt := st.chat_input("Example: tighten indemnity language and simplify definitions"):
            if "create_chat_messages" not in st.session_state:
                st.session_state.create_chat_messages = []
            st.session_state.create_chat_messages.append({"role": "user", "content": prompt})

        if st.session_state.get("create_chat_messages") and st.session_state.create_chat_messages[-1]["role"] == "user":
            pending_prompt = st.session_state.create_chat_messages[-1]["content"]
            with st.chat_message("assistant"):
                with st.spinner("Applying revisions..."):
                    llm = get_llm()
                    messages = [
                        {
                            "role": "system",
                            "content": "You are a legal document editing assistant. Apply requested changes and return the FULL updated document."
                        },
                        {
                            "role": "user",
                            "content": (
                                f"Draft objective: {st.session_state.get('create_setup_goal')}\n\n"
                                f"Current document:\n\n{st.session_state.create_generated_text}\n\n"
                                f"Revision request: {pending_prompt}\n\n"
                                "Return only the full updated document."
                            )
                        }
                    ]
                    try:
                        response = llm.chat(messages, temperature=0.2, max_tokens=4096)
                        st.session_state.create_generated_text = response
                        st.session_state.create_chat_messages.append({"role": "assistant", "content": "Revision applied."})
                        st.rerun()
                    except Exception as e:
                        err = f"Error: {e}"
                        st.error(err)
                        st.session_state.create_chat_messages.append({"role": "assistant", "content": err})


# ======================================================================
# PATH C: Learn from a Document Workflow
# ======================================================================

def render_learn_workflow():
    """Learn from uploaded documents and maintain a personal style library."""

    if st.button("Back to Home", key="back_from_learn"):
        st.session_state.workflow_mode = None
        st.rerun()

    render_workflow_header(
        "Learn from a Document",
        "Upload reference examples, index them, and keep a personal style library for drafting/editing.",
        progress=1.0,
        step_note="Single step: Upload, index, and review style sources",
    )

    doc_type = st.selectbox(
        "Learning category",
        options=["reference"] + list(DOCUMENT_TYPES.keys()),
        format_func=lambda x: "Reference (general)" if x == "reference" else DOCUMENT_TYPES[x]["label"],
        key="learn_doc_type",
    )

    uploaded_files = st.file_uploader(
        "Upload one or more style documents",
        type=["pdf", "doc", "docx", "docm", "txt"],
        accept_multiple_files=True,
        key="learn_upload_docs",
        help="These files are saved to your personal style library and can also be indexed for retrieval.",
    )

    index_for_retrieval = st.checkbox(
        "Also index these files for retrieval",
        value=True,
        key="learn_index_for_retrieval",
        help="Keeps these documents searchable in document generation and Q&A retrieval.",
    )

    if uploaded_files:
        st.caption(f"Ready to process: {len(uploaded_files)} file(s) under '{doc_type}'.")
    else:
        st.caption("Upload one or more files to add personal style sources.")

    if uploaded_files and st.button("Save to Personal Style Library", type="primary", use_container_width=True, key="learn_process_btn"):
        with st.spinner("Saving style documents..."):
            summary = save_style_library_documents(
                uploaded_files,
                document_type=doc_type,
                index_for_retrieval=index_for_retrieval,
            )
        st.success("Personal style library updated.")
        st.caption(
            f"Saved: {summary['saved']} | Failed: {summary['failed']}"
        )
        if index_for_retrieval:
            idx = summary.get("indexed", {})
            st.caption(
                f"Indexed: {idx.get('new_files', 0)} | Chunks: {idx.get('total_chunks', 0)} | "
                f"Skipped duplicates: {idx.get('skipped_files', 0)} | Failed: {idx.get('failed_files', 0)}"
            )
        st.rerun()

    st.divider()
    st.markdown("### Personal Style Library")
    style_docs = get_personal_style_documents()
    if not style_docs:
        st.info("No personal style documents yet. Upload documents above to build your style library.")
        return

    selected_style_id = st.selectbox(
        "Saved style documents",
        options=[doc["id"] for doc in style_docs],
        format_func=lambda doc_id: next(
            (
                f"{d['name']} | {d.get('document_type', 'reference')} | {d.get('char_count', 0):,} chars"
                for d in style_docs if d["id"] == doc_id
            ),
            str(doc_id),
        ),
        key="learn_style_doc_selector",
    )

    selected_doc = get_personal_style_document(int(selected_style_id)) if selected_style_id else None
    if selected_doc:
        st.caption(
            f"Updated: {selected_doc.get('updated_at', '')}"
        )
        st.text_area(
            "Style document preview",
            value=selected_doc.get("content_text", ""),
            height=320,
            disabled=False,
            key=f"learn_style_preview_{selected_style_id}",
        )


# ======================================================================
# Settings Page (Accessed from Sidebar)
# ======================================================================

def render_settings_page():
    """Render the full settings page."""

    if st.button("Back to Workspace", key="back_from_settings"):
        st.session_state.show_settings = False
        st.rerun()

    render_workflow_header(
        "Settings",
        "Manage AI provider, profile, and admin controls.",
        step_note="Changes apply to your current session immediately after saving.",
    )

    current_user = st.session_state.current_user
    st.caption(f"Signed in as: {current_user.full_name} ({current_user.role})")

    settings_tab1, settings_tab2, settings_tab3 = st.tabs([
        "LLM Provider",
        "Profile",
        "Admin" if current_user.is_admin() else "Info",
    ])

    with settings_tab1:
        st.markdown("""
        <div class="info-card">
            The AI language model powers document generation and editing. Choose your preferred provider.
        </div>
        """, unsafe_allow_html=True)

        if FORCE_OLLAMA_FOR_ALL_USERS:
            provider = "ollama"
            st.info("Self-host mode: Ollama is enforced for all users.")
        else:
            provider_options = ["ollama", "hf_local", "openai"]
            current_provider = st.session_state.llm_provider if st.session_state.llm_provider in provider_options else "ollama"
            provider = st.radio(
                "Select AI Provider",
                options=provider_options,
                format_func=lambda p: "Ollama (Local, free)" if p == "ollama" else ("Hugging Face Local (folder path)" if p == "hf_local" else "OpenAI (Cloud API)"),
                index=provider_options.index(current_provider),
                key="settings_provider",
            )

        model = st.session_state.llm_model
        base_url = st.session_state.llm_base_url
        api_key = st.session_state.openai_api_key
        hf_local_model_path = str(st.session_state.get("hf_local_model_path", "") or "")
        hf_local_max_new_tokens_raw = st.session_state.get("hf_local_max_new_tokens", "2048")
        try:
            hf_local_max_new_tokens = int(hf_local_max_new_tokens_raw)
        except Exception:
            hf_local_max_new_tokens = 2048
        is_openai_key_valid = True

        if provider == "ollama":
            st.markdown("#### Ollama Configuration")
            base_url = st.text_input(
                "Ollama API URL",
                value=st.session_state.llm_base_url,
                key="settings_base_url",
                help="Default: http://localhost:11434/v1",
            )
            tmp_llm = LLMBackend(provider="ollama", base_url=base_url)
            models = tmp_llm.list_models()
            if models:
                model = st.selectbox(
                    "Select Model",
                    options=models,
                    index=models.index(st.session_state.llm_model)
                    if st.session_state.llm_model in models
                    else 0,
                    key="settings_model_select",
                )
                st.success(f"Connected to Ollama - {len(models)} model(s) available")
            else:
                model = st.text_input(
                    "Model name",
                    value=st.session_state.llm_model,
                    key="settings_model_text",
                    placeholder="llama3.1:8b",
                )
                st.error("Could not connect to Ollama. Make sure it's running.")
                st.caption("Start Ollama: `ollama serve`")

        elif provider == "hf_local":
            st.markdown("#### Hugging Face Local Configuration")
            hf_local_model_path = st.text_input(
                "Local model folder path",
                value=hf_local_model_path,
                key="settings_hf_model_path",
                placeholder=r"C:\models\my-hf-model",
                help="Folder must contain config.json, model weights, and tokenizer files.",
            )
            hf_local_max_new_tokens = st.number_input(
                "Max new tokens per generation",
                min_value=64,
                max_value=8192,
                step=64,
                value=hf_local_max_new_tokens,
                key="settings_hf_max_tokens",
            )
            model = os.path.basename(hf_local_model_path.rstrip("\\/")) if hf_local_model_path.strip() else "hf-local"
            base_url = "local://hf"

            if hf_local_model_path.strip():
                if os.path.isdir(hf_local_model_path.strip()):
                    st.success("Model path exists. Save settings, then generate from the workspace.")
                else:
                    st.warning("Model path does not exist on this machine.")
            else:
                st.info("Enter a local Hugging Face model folder path.")

        else:
            st.markdown("#### OpenAI Configuration")
            api_key = st.text_input(
                "OpenAI API Key",
                value=st.session_state.openai_api_key if st.session_state.openai_api_key not in ["", "your-api-key-here"] else "",
                type="password",
                key="settings_openai_key",
                placeholder="sk-proj-...",
            )
            model_options = ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo"]
            current_model = st.session_state.llm_model if st.session_state.llm_provider == "openai" else "gpt-4o-mini"
            model = st.selectbox(
                "Model",
                options=model_options,
                index=model_options.index(current_model) if current_model in model_options else 0,
                key="settings_openai_model",
            )
            base_url = "https://api.openai.com/v1"

            is_openai_key_valid = bool(api_key and api_key.startswith("sk-") and len(api_key) > 20)
            if is_openai_key_valid:
                st.success("API key format looks valid")
            elif api_key:
                st.warning("API key format may be invalid (should start with 'sk-')")
            else:
                st.info("Enter your OpenAI API key to enable cloud generation.")

        save_disabled = (provider == "openai" and not is_openai_key_valid) or (provider == "hf_local" and not hf_local_model_path.strip())
        if provider == "openai" and save_disabled:
            st.caption("Enter a valid OpenAI API key to save OpenAI settings.")
        if provider == "hf_local" and save_disabled:
            st.caption("Enter a valid local model folder path to save HF Local settings.")

        if st.button("Save Settings", type="primary", use_container_width=True, disabled=save_disabled):
            st.session_state.llm_provider = provider
            st.session_state.llm_model = model
            st.session_state.llm_base_url = base_url
            if provider == "openai":
                st.session_state.openai_api_key = api_key
            if provider == "hf_local":
                st.session_state.hf_local_model_path = hf_local_model_path.strip()
                st.session_state.hf_local_max_new_tokens = str(int(hf_local_max_new_tokens))
                st.session_state.hf_local_error = ""
            st.toast("Settings saved")
            st.success("AI provider settings saved successfully.")
            st.rerun()

    with settings_tab2:
        render_profile_settings(auth_manager, current_user)

    with settings_tab3:
        if current_user.is_admin():
            render_admin_panel(auth_manager, current_user)
        else:
            st.markdown("""
            <div class="info-card">
                <strong>Account Information</strong><br><br>
                You are signed in as a standard user.<br>
                Contact an administrator for role changes or workspace-level requests.
            </div>
            """, unsafe_allow_html=True)


# ======================================================================
# Knowledge Base Page
# ======================================================================

def render_knowledge_base_page():
    """Render the Knowledge Base management page."""

    current_user = st.session_state.current_user
    if not current_user.is_admin():
        st.error("Access denied. Knowledge Base management is admin-only.")
        return

    # Back button
    if st.button("Back to Workspace", key="back_from_kb"):
        st.session_state.show_knowledge_base = False
        st.rerun()

    render_workflow_header(
        "Knowledge Base",
        "Manage learned templates, indexed files, scanner settings, and scan history.",
        step_note="Admin only: use this view to keep retrieval quality high.",
    )

    knowledge_db = get_knowledge_db()
    # Lazy-init scanner only on Knowledge Base page (admin/manual use).
    scope_user_id = get_data_scope_user_id()
    scanner = ScannerManager.get_instance(scope_user_id)
    if scanner is None:
        try:
            scanner = get_background_scanner(scope_user_id, get_llm(), get_collection(), knowledge_db)
        except Exception as e:
            logger.error(f"Failed to initialize scanner on KB page: {e}")
            scanner = None

    stats = knowledge_db.get_stats(user_id=scope_user_id)

    # Overview stats
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("Documents Indexed", stats["total_documents"])

    with col2:
        st.metric("Document Types", stats["document_types"])

    with col3:
        st.metric("Total Scans", stats["total_scans"])

    with col4:
        if stats["last_scan"]:
            try:
                import dateutil.parser
                import datetime
                last_scan = dateutil.parser.parse(stats["last_scan"])
                time_ago = datetime.datetime.now(datetime.timezone.utc) - last_scan
                hours_ago = int(time_ago.total_seconds() / 3600)
                st.metric("Last Scan", f"{hours_ago}h ago")
            except:
                st.metric("Last Scan", "Recently")
        else:
            st.metric("Last Scan", "Never")

    st.divider()

    # Tabs for different sections
    kb_tab1, kb_tab2, kb_tab3, kb_tab4 = st.tabs([
        "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ Learned Templates",
        "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Indexed Documents",
        "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¾ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Scanner Settings",
        "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â  Scan History"
    ])

    # Tab 1: Learned Templates
    with kb_tab1:
        st.markdown("### Learned Document Templates")
        st.caption("Templates automatically learned from your real documents")

        doc_types = knowledge_db.get_all_document_types(user_id=scope_user_id)

        if not doc_types:
            st.info("No learned templates yet. Run a scan to build templates from your documents.")
        else:
            for type_info in doc_types:
                doc_type = type_info["type"]
                count = type_info["count"]

                with st.expander(f"ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦ÃƒÂ¢Ã¢â€šÂ¬Ã…â€œÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Â¦Ãƒâ€šÃ‚Â¾ {doc_type} ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Learned from {count} documents"):
                    learned = knowledge_db.get_learned_template(doc_type, user_id=scope_user_id)

                    if learned:
                        st.caption(f"Last updated: {learned['last_updated']}")

                        st.markdown("**Learned Fields:**")
                        for field in learned["fields"]:
                            confidence = field.get("confidence", 0)
                            frequency = field.get("frequency", 0)
                            bar_color = "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢" if confidence >= 0.7 else "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡" if confidence >= 0.5 else "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â°ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â´"
                            st.markdown(
                                f"{bar_color} **{field['label']}** ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã‚Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â "
                                f"{confidence*100:.0f}% confidence ({frequency}/{count} documents)"
                            )

    # Tab 2: Indexed Documents
    with kb_tab2:
        st.markdown("### Indexed Documents by Type")

        if not doc_types:
            st.info("No documents indexed yet.")
        else:
            selected_type = st.selectbox(
                "Select document type to view:",
                options=[t["type"] for t in doc_types],
                key="kb_doc_type_select"
            )

            if selected_type:
                files = knowledge_db.get_scanned_files_by_type(selected_type, user_id=scope_user_id)

                st.caption(f"Found {len(files)} documents")

                for file_info in files[:50]:  # Limit display to 50
                    file_name = os.path.basename(file_info["file_path"])
                    file_size_kb = file_info["file_size"] / 1024

                    col_a, col_b, col_c = st.columns([3, 1, 1])
                    with col_a:
                        st.markdown(f"**{file_name}**")
                    with col_b:
                        st.caption(f"{file_size_kb:.1f} KB")
                    with col_c:
                        st.caption(file_info["scan_date"][:10])

    # Tab 3: Scanner Settings
    with kb_tab3:
        st.markdown("### Scanner Configuration")

        scope_user_id = get_data_scope_user_id()
        scanner = ScannerManager.get_instance(scope_user_id)
        if scanner:
            status = scanner.get_status()

            st.markdown("#### Status")
            if status["is_running"]:
                st.success("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦ Scanner is running")

                if status.get("current_progress"):
                    prog = status["current_progress"]
                    st.progress(prog["current"] / max(prog["total"], 1))
                    st.caption(f"Processing: {prog['file']}")
                    st.caption(f"{prog['current']} / {prog['total']} files")
            else:
                st.warning("ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¡ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¯ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¸ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â Scanner is stopped")

            st.divider()

            st.markdown("#### Scan Paths")
            scan_paths = status.get("scan_paths", [])
            for path in scan_paths:
                exists = os.path.exists(path)
                icon = "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦" if exists else "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬ÃƒÂ¢Ã¢â‚¬Å¾Ã‚Â¢"
                st.markdown(f"{icon} `{path}`")

            st.divider()
            st.markdown("#### Manual Controls")

            col_a, col_b = st.columns(2)

            with col_a:
                if st.button("Run Scan Now", type="primary", use_container_width=True):
                    with st.spinner("Running scan..."):
                        scanner.run_scan_once()
                    st.success("Scan completed")
                    st.rerun()

            with col_b:
                if st.button("Rebuild Templates", use_container_width=True):
                    with st.spinner("Rebuilding learned templates..."):
                        llm = get_llm()
                        collection = get_collection()
                        scanner_instance = DocumentScanner(
                            llm=llm,
                            chroma_collection=collection,
                            knowledge_db=knowledge_db,
                            user_id=scope_user_id,
                        )
                        scanner_instance.build_learned_templates()
                        st.success("Templates rebuilt")
                        st.rerun()

            st.divider()

            st.markdown("#### Advanced Settings")

            new_interval = st.number_input(
                "Scan interval (minutes)",
                min_value=5,
                max_value=1440,
                value=status.get("scan_interval_minutes", 30),
                key="kb_scan_interval"
            )

            if st.button("Update Interval", key="update_scan_interval"):
                # This would require updating the scanner config
                st.info("Restart the scanner for the new interval to take effect")

        else:
            st.warning("Background scanner is not initialized yet.")
            if st.button("Initialize Scanner", key="kb_init_scanner", use_container_width=True):
                try:
                    _ = get_background_scanner(scope_user_id, get_llm(), get_collection(), knowledge_db)
                    st.success("Scanner initialized")
                    st.rerun()
                except Exception as init_err:
                    st.error(f"Failed to initialize scanner: {init_err}")

    # Tab 4: Scan History
    with kb_tab4:
        st.markdown("### Recent Scan History")

        history = knowledge_db.get_scan_history(limit=20, user_id=scope_user_id)

        if not history:
            st.info("No scan history yet.")
        else:
            for scan in history:
                status_icon = "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¦ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Ã¢â‚¬Å“ÃƒÆ’Ã†â€™Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡Ãƒâ€šÃ‚Â¬ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¦" if scan["status"] == "completed" else "ÃƒÆ’Ã†â€™Ãƒâ€ Ã¢â‚¬â„¢ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â¢ÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚ÂÃƒÆ’Ã†â€™ÃƒÂ¢Ã¢â€šÂ¬Ã…Â¡ÃƒÆ’Ã¢â‚¬Å¡Ãƒâ€šÃ‚Â³"

                with st.expander(f"{status_icon} Scan on {scan['scan_start'][:10]}"):
                    col_a, col_b, col_c = st.columns(3)

                    with col_a:
                        st.metric("Scanned", scan["files_scanned"])
                    with col_b:
                        st.metric("Indexed", scan["files_indexed"])
                    with col_c:
                        st.metric("Failed", scan["files_failed"])

                    if scan["scan_end"]:
                        try:
                            import dateutil.parser
                            start = dateutil.parser.parse(scan["scan_start"])
                            end = dateutil.parser.parse(scan["scan_end"])
                            duration = (end - start).total_seconds()
                            st.caption(f"Duration: {duration:.0f} seconds")
                        except:
                            pass

                    if scan.get("errors"):
                        st.markdown("**Errors:**")
                        st.text(scan["errors"])


# ======================================================================
# Main Application
# ======================================================================

def main():
    """Main application entry point."""

    # Check authentication status
    if not require_auth(st.session_state):
        # Show verification, registration, or login page
        if st.session_state.get("verify_email"):
            render_verification_page(auth_manager)
        elif st.session_state.get("show_register", False):
            render_registration_page(auth_manager)
        else:
            render_login_page(auth_manager)
        return

    # User is authenticated
    current_user = st.session_state.current_user

    if current_user.must_change_password:
        st.warning("Password update required. You can continue, but update it from Settings -> Profile.")

    # Auto-connect all users to the host's Ollama instance (no onboarding needed)
    if not st.session_state.get("onboarding_complete", False):
        if st.session_state.llm_provider == "ollama":
            st.session_state.onboarding_complete = True
        elif st.session_state.llm_provider == "openai" and not st.session_state.openai_api_key:
            render_onboarding_wizard()
            return
        else:
            st.session_state.onboarding_complete = True

    st.title("Corporate Law Document Generator")
    st.caption("Document generation workspace")

    top_actions = st.columns(5)
    with top_actions[0]:
        if st.button("Workspace Home", use_container_width=True, key="top_workspace"):
            st.session_state.workflow_mode = None
            st.session_state.show_settings = False
            st.session_state.show_knowledge_base = False
            st.session_state.show_model_improvement = False
            st.rerun()
    with top_actions[1]:
        if st.button("Settings", use_container_width=True, key="top_settings"):
            st.session_state.show_settings = True
            st.session_state.show_knowledge_base = False
            st.session_state.show_model_improvement = False
            st.rerun()
    with top_actions[2]:
        if current_user.is_admin():
            if st.button("Knowledge Base", use_container_width=True, key="top_knowledge"):
                st.session_state.show_knowledge_base = True
                st.session_state.show_settings = False
                st.session_state.show_model_improvement = False
                st.rerun()
        else:
            st.button("Knowledge Base (Admin)", use_container_width=True, key="top_knowledge_disabled", disabled=True)
    with top_actions[3]:
        if current_user.is_admin():
            if st.button("Model Improvement", use_container_width=True, key="top_model_improvement"):
                st.session_state.show_model_improvement = True
                st.session_state.show_knowledge_base = False
                st.session_state.show_settings = False
                st.rerun()
        else:
            st.button("Model Improvement (Admin)", use_container_width=True, key="top_model_improvement_disabled", disabled=True)
    with top_actions[4]:
        if st.button("Sign Out", use_container_width=True, key="top_logout"):
            logout(st.session_state)
            st.rerun()

    current_view = (
        "Settings" if st.session_state.get("show_settings")
        else "Knowledge Base" if st.session_state.get("show_knowledge_base")
        else "Model Improvement" if st.session_state.get("show_model_improvement")
        else "Workspace"
    )
    st.caption(f"Current view: {current_view}")

    llm = get_llm()
    if not llm.is_available():
        if st.session_state.llm_provider == "ollama":
            st.markdown("""
            <div class="error-card">
                <strong>Ollama Not Running</strong><br><br>
                Start Ollama (<code>ollama serve</code>) and refresh this page, or switch providers in Settings.
            </div>
            """, unsafe_allow_html=True)
        elif st.session_state.llm_provider == "hf_local":
            reason = st.session_state.get("hf_local_error", "HF Local model is not available. Check Settings.")
            st.markdown(f"""
            <div class="error-card">
                <strong>HF Local Model Unavailable</strong><br><br>
                {html_lib.escape(reason)}
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="error-card">
                <strong>OpenAI API Key Required</strong><br><br>
                Add a valid API key in Settings before generating documents.
            </div>
            """, unsafe_allow_html=True)

    if st.session_state.get("show_settings"):
        render_settings_page()
        return

    if st.session_state.get("show_knowledge_base"):
        if current_user.is_admin():
            render_knowledge_base_page()
        else:
            st.error("Access denied. Knowledge Base management is admin-only.")
            st.session_state.show_knowledge_base = False
        return

    if st.session_state.get("show_model_improvement"):
        if current_user.is_admin():
            render_model_improvement_page()
        else:
            st.error("Access denied. Model improvement controls are admin-only.")
            st.session_state.show_model_improvement = False
        return

    workflow_mode = st.session_state.get("workflow_mode")

    if workflow_mode == "edit":
        render_edit_workflow()
    elif workflow_mode == "create":
        render_create_workflow()
    elif workflow_mode == "learn":
        render_learn_workflow()
    else:
        render_landing_page()

    render_footer()


if __name__ == "__main__":
    main()










